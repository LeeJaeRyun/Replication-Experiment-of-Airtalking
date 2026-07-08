from __future__ import annotations

import csv
import json
from collections import Counter
from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
CITY_DIR = ROOT / "outputs" / "cityscapes_semantic_measurement"
ORIGINAL_DIR = ROOT / "outputs" / "airtalking_cityscapes_feature_paperhw"
RUN_DIR = ROOT / "outputs" / "airtalking_cityscapes_calibrated_final_p012"
FIG_DIR = RUN_DIR / "figures"

CITY_SUMMARY_JSON = CITY_DIR / "cityscapes_semantic_summary.json"
SUMMARY_CSV = RUN_DIR / "summary_metrics.csv"
METADATA_JSON = RUN_DIR / "run_metadata.json"
VERIFY_CSV = RUN_DIR / "verification_against_paper_calibrated_final_p012.csv"
ORIGINAL_VERIFY_CSV = ORIGINAL_DIR / "verification_against_paper_cityscapes_feature_paperhw.csv"
REPORT_PATH = RUN_DIR / "AirTalking_easy_reproduction_report_KR.docx"

LATIN_FONT = "Calibri"
KOREAN_FONT = "Malgun Gothic"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "555555"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
POSITIVE_FILL = "EEF7F1"
BORDER = "D7DBE2"


def require_inputs() -> None:
    required = [
        CITY_SUMMARY_JSON,
        SUMMARY_CSV,
        METADATA_JSON,
        VERIFY_CSV,
        ORIGINAL_VERIFY_CSV,
        FIG_DIR / "finished_requests.png",
        FIG_DIR / "semantic_vs_nonsemantic_300m.png",
        FIG_DIR / "average_time_cost.png",
    ]
    missing = [str(path.relative_to(ROOT)) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required files:\n" + "\n".join(missing))


def set_run_font(
    run,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = LATIN_FONT
    style._element.get_or_add_rPr()
    rfonts = style._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        style._element.rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def paragraph_border_bottom(paragraph, color: str = BORDER, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(table, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_first_row_as_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_cell_text(cell, text: str, size: float = 8.8, bold: bool = False, color: str = "000000", align=None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.12
    if align is not None:
        para.alignment = align
    run = para.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], font_size: float = 8.6) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    mark_first_row_as_header(table)
    set_cell_margins(table)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, header, size=font_size, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1) and len(values) > 2 else None
            set_cell_text(cells[idx], value, size=font_size, align=align)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    mark_first_row_as_header(table)
    set_cell_margins(table, top=150, bottom=150, start=190, end=190)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(title)
    set_run_font(r_title, size=10.5, color=NAVY, bold=True)
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    p_body.paragraph_format.line_spacing = 1.12
    r_body = p_body.add_run(body)
    set_run_font(r_body, size=9.8, color="222222")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_para(doc: Document, text: str, size: float = 10.6, color: str = "000000", bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.4)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=8.7, color=MUTED, italic=True)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.1) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image = p.add_run().add_picture(str(path), width=Inches(width))
    image._inline.docPr.set("descr", caption)
    image._inline.docPr.set("title", caption.split(".", 1)[0])
    add_caption(doc, caption)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    set_style_font(styles["Heading 1"], size=16, color=BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Heading 2"], size=13, color=BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    set_style_font(styles["List Bullet"], size=10.4)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Bullet"].paragraph_format.space_after = Pt(8)
    styles["List Bullet"].paragraph_format.line_spacing = 1.167

    header = section.header
    header.paragraphs[0].text = ""
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    hr = hp.add_run("AirTalking reproduction | Cityscapes-based approximate run")
    set_run_font(hr, size=8.5, color=MUTED)
    paragraph_border_bottom(hp, color=BORDER, size="6")

    footer = section.footer
    footer.paragraphs[0].text = ""
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Generated: " + date.today().isoformat())
    set_run_font(fr, size=8.5, color=MUTED)


def add_masthead(doc: Document, metadata: dict, city: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("AirTalking 논문 재현 실험 보고서")
    set_run_font(r, size=23, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Cityscapes 공개 데이터 기반 근사 재현 및 보정 실험")
    set_run_font(r, size=13.5, color=MUTED)

    rows = [
        ["사용 데이터", f"Cityscapes train/val {city['num_samples']:,}쌍"],
        ["핵심 측정값", f"semantic compression ratio rho_c={city['rho_c_feature_uncompressed_mean']:.6f}"],
        ["최종 실행", f"{metadata['paper_params']['repeats']}회 반복, seed={metadata['assumed_params']['seed']}"],
        ["보고서 성격", "공식 코드가 없는 논문의 공개 정보 기반 approximate reproduction"],
    ]
    add_table(doc, ["항목", "내용"], rows, [2200, 7160], font_size=9.0)

    add_callout(
        doc,
        "한 줄 결론",
        "공개 데이터로 확인 가능한 semantic payload 크기 비율은 논문값과 거의 같게 재현됐다. "
        "시뮬레이션 결과는 보정 후 개선됐지만, 논문 원 그래프와 완전히 같지는 않다. "
        "차이는 주로 논문이 공개하지 않은 request 생성, workload, 전력, 간섭 처리, 정책 hyperparameter에서 나온다.",
        fill=POSITIVE_FILL,
    )


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def count_verdicts(rows: list[dict[str, str]]) -> Counter:
    return Counter(row["verdict"].strip().lower() for row in rows)


def avg_capped_relative_error(rows: list[dict[str, str]], cap: float = 5.0) -> float:
    values = []
    for row in rows:
        try:
            values.append(min(float(row["relative_error"]), cap))
        except (TypeError, ValueError):
            continue
    return sum(values) / len(values)


def metric_row(rows: list[dict[str, str]], mode: str, area: int, policy: str) -> dict[str, str]:
    for row in rows:
        if row["mode"] == mode and int(row["area"]) == area and row["policy"] == policy:
            return row
    raise KeyError((mode, area, policy))


def fmt_num(value: float | str, digits: int = 1) -> str:
    value = float(value)
    return f"{value:.{digits}f}"


def write_problem_definition(doc: Document) -> None:
    doc.add_heading("1. 문제정의", level=1)
    add_para(
        doc,
        "AirTalking 논문은 여러 UAV가 움직이는 환경에서 장치의 데이터를 다른 UAV 또는 장치로 보내는 문제를 다룬다. "
        "원본 영상 데이터를 그대로 보내면 전송해야 할 bit 수가 커지고, 전송 시간과 비행 에너지까지 같이 증가한다.",
    )
    add_para(
        doc,
        "논문의 핵심 질문은 간단하다. 영상 전체를 보내는 대신, 영상에서 의미 정보만 추출해서 보내면 "
        "같은 시간 안에 더 많은 요청을 처리하고 에너지를 줄일 수 있는가?",
    )
    add_bullet(doc, "입력 데이터: 도시 주행 이미지와 semantic segmentation label")
    add_bullet(doc, "통신 대상: UAV, 지상 장치, UAV 간 링크")
    add_bullet(doc, "비교 대상: semantic processing 사용 방식과 원본 데이터 전송 방식")
    add_bullet(doc, "평가 지표: 처리 완료 요청 수, 평균 처리 시간, 비행/비비행 에너지, 이동 거리")
    add_callout(
        doc,
        "이 재현의 범위",
        "딥러닝 모델을 새로 학습한 것이 아니라, 공개 Cityscapes 데이터에서 semantic payload 크기 비율을 측정하고 "
        "그 값을 UAV 통신 시뮬레이터에 넣어 논문 그래프를 근사했다.",
        fill=CALLOUT_FILL,
    )


def write_proposal(doc: Document) -> None:
    doc.add_heading("2. 제안내용", level=1)
    add_para(
        doc,
        "논문이 제안한 방식은 semantic communication이다. 원본 이미지를 전부 보내지 않고, "
        "semantic encoder가 장면의 의미 정보를 먼저 뽑아 작은 payload로 만들고, receiver 쪽에서 decoder가 이를 복원하거나 활용한다.",
    )
    add_para(
        doc,
        "논문에는 semantic encoder/decoder 관련 단서가 있다. modified U-Net과 Pix2PixHD 계열 구조, "
        "Cityscapes 기반 semantic segmentation 모델, Jetson 계열 장치에서의 encoder/decoder 처리율이 언급된다. "
        "다만 공식 weight, 학습 코드, feature tensor 정의, 정확한 압축 bitstream은 공개되어 있지 않다.",
    )
    add_table(
        doc,
        ["구성요소", "역할", "이번 재현에서 사용한 처리"],
        [
            ["Semantic encoder", "이미지에서 의미 feature 추출", "논문 공개 처리율 91.3 Mbps 사용"],
            ["Semantic decoder", "전송된 의미 feature 복원/활용", "논문 공개 처리율 23.23 Mbps 사용"],
            ["rho_c", "원본 대비 semantic payload 비율", "Cityscapes label 기반으로 직접 측정"],
            ["rho_r", "decoder 후 복원 비율 proxy", "논문 Table III 값 3 사용"],
            ["정책 알고리즘", "어떤 UAV/링크/처리를 선택할지 결정", "Stochastic, LinUCB, SA, Greedy, MCTS 구현"],
        ],
        [1750, 2700, 4910],
        font_size=8.4,
    )
    add_para(
        doc,
        "기술적으로는 컴퓨터비전 데이터셋, 딥러닝 기반 semantic segmentation 개념, Python/NumPy 시뮬레이션, "
        "무선통신의 path loss/SINR/Shannon capacity 계산, Matplotlib 기반 결과 시각화를 조합했다.",
    )


def write_experiment(doc: Document, metadata: dict, city: dict) -> None:
    doc.add_heading("3. 실험내용", level=1)
    doc.add_heading("3.1 사용한 데이터", level=2)
    add_table(
        doc,
        ["데이터", "실제 사용 여부", "설명"],
        [
            ["Cityscapes leftImg8bit train/val", "사용", "RGB 원본 이미지. raw payload 기준으로 사용"],
            ["Cityscapes gtFine train/val labelIds", "사용", "semantic label map. 의미 feature 크기 측정에 사용"],
            ["Cityscapes test", "사용 안 함", "정답 label이 dummy/ignore라 rho_c 측정에 넣지 않음"],
            ["trainextra/coarse/right/disparity/camera", "사용 안 함", "이번 실험 목적에 필요하지 않아 입력 데이터로 넣지 않음"],
            ["논문 원 실험 데이터/코드", "없음", "공식 코드와 raw simulation result가 공개되어 있지 않음"],
        ],
        [2450, 1600, 5310],
        font_size=8.4,
    )
    add_para(
        doc,
        f"실제로 계산에 들어간 Cityscapes sample은 train/val 총 {city['num_samples']:,}쌍이다. "
        "이전 보고서에서 test split이나 trainextra를 '미사용'이라고 쓴 것은 다운로드 항목을 구분하려는 표현이었고, "
        "이번 보고서에서는 애초에 입력 데이터로 넣은 두 항목만 명확히 표시했다.",
    )

    doc.add_heading("3.2 semantic payload 측정", level=2)
    raw_mb = city["raw_uncompressed_bytes_mean"] / 1024 / 1024
    sem_mb = city["semantic_feature_bytes_mean"] / 1024 / 1024
    add_table(
        doc,
        ["측정 항목", "값", "의미"],
        [
            ["원본 RGB 크기", f"{raw_mb:.2f} MB", "2048 x 1024 x 3 byte 이미지"],
            ["semantic feature proxy", f"{sem_mb:.2f} MB", "labelIds를 feature scale 0.56으로 downsample한 tensor"],
            ["rho_c 측정값", f"{city['rho_c_feature_uncompressed_mean']:.6f}", "논문값 0.104와 거의 같음"],
            ["rho_r", f"{city['rho_r_proxy']:.1f}", "논문 Table III 값 사용"],
        ],
        [2350, 1700, 5310],
        font_size=8.6,
    )
    add_para(
        doc,
        "여기서 rho_c는 원본 payload 대비 semantic payload가 얼마나 작은지를 뜻한다. "
        "측정값 0.104464는 논문에 적힌 0.104와 매우 가까워서, 공개 데이터로 검증 가능한 부분은 잘 맞는다.",
    )

    doc.add_heading("3.3 시뮬레이션 설정", level=2)
    paper = metadata["paper_params"]
    assumed = metadata["assumed_params"]
    add_table(
        doc,
        ["분류", "사용 값", "출처/이유"],
        [
            ["논문 공개값", f"UAV {paper['n_uav']}대, device {paper['n_device']}개, {paper['t_slots']} slots", "논문 Table III 기반"],
            ["통신값", "80 MHz, 5 GHz, UAV Tx 0.2 W, device Tx 0.1 W", "논문 공개값"],
            ["semantic 처리율", "encoder 91.3 Mbps, decoder 23.23 Mbps", "논문 공개값"],
            ["request probability", f"{assumed['request_probability']}", "논문 미공개. 처리 요청 수가 논문 그래프 범위에 들어오도록 보정"],
            ["workload", f"평균 {assumed['workload_mean_bits']/1e6:.0f} Mb", "논문 미공개. 평균 처리 시간 scale을 맞추기 위한 clipped normal workload"],
            ["전력", f"P_move {assumed['p_move']:.0f} W, P_hover {assumed['p_hover']:.0f} W", "논문 미공개. per-task effective simulator power로 보정"],
            ["정책 탐색 budget", f"LinUCB 후보 {assumed['linucb_candidate_samples']}, SA iter {assumed['sa_iterations']}", "논문 미공개. 반복 가능한 계산량으로 제한"],
        ],
        [2100, 3300, 3960],
        font_size=8.1,
    )
    add_callout(
        doc,
        "임의로 숨긴 값은 없음",
        "논문에 없는 값은 코드의 assumed_params로 분리했고, 보고서에도 따로 표시했다. "
        "논문값처럼 말하지 않고, 왜 그런 값을 썼는지 설명한다. 값은 결과를 맞추기 위한 보정값이지 논문 저자의 실제 설정이라고 주장할 수 없다.",
        fill=CAUTION_FILL,
    )


def write_results(doc: Document, summary_rows: list[dict[str, str]], original_verify: list[dict[str, str]], final_verify: list[dict[str, str]]) -> None:
    doc.add_heading("4. 실험결과", level=1)
    original_counts = count_verdicts(original_verify)
    final_counts = count_verdicts(final_verify)
    original_err = avg_capped_relative_error(original_verify)
    final_err = avg_capped_relative_error(final_verify)
    improvement = (original_err - final_err) / original_err * 100.0

    add_table(
        doc,
        ["구분", "Match", "Partial", "Mismatch", "평균 capped relative error"],
        [
            [
                "보정 전",
                str(original_counts.get("match", 0)),
                str(original_counts.get("partial", 0)),
                str(original_counts.get("mismatch", 0)),
                f"{original_err:.3f}",
            ],
            [
                "보정 후",
                str(final_counts.get("match", 0)),
                str(final_counts.get("partial", 0)),
                str(final_counts.get("mismatch", 0)),
                f"{final_err:.3f}",
            ],
        ],
        [1900, 1400, 1400, 1700, 2960],
        font_size=8.6,
    )
    add_para(
        doc,
        f"보정 후 Match는 {final_counts.get('match', 0)}개로 늘었고, 평균 capped relative error는 "
        f"{original_err:.3f}에서 {final_err:.3f}로 약 {improvement:.1f}% 낮아졌다. "
        "즉 결과가 개선됐지만, 논문 그래프와 거의 동일하다고 말할 정도는 아니다.",
    )

    doc.add_heading("4.1 완료 요청 수", level=2)
    policies = ["Stochastic", "LinUCB", "SA", "Greedy", "MCTS"]
    rows = []
    for area in [100, 200, 300, 400, 500]:
        row = [f"{area} x {area}"]
        for policy in policies:
            row.append(fmt_num(metric_row(summary_rows, "semantic", area, policy)["finished"], 1))
        rows.append(row)
    add_table(doc, ["Area", *policies], rows, [1500, 1550, 1550, 1550, 1550, 1660], font_size=8.0)
    add_figure(
        doc,
        FIG_DIR / "finished_requests.png",
        "그림 1. 면적별 완료 요청 수. 대부분의 정책에서 공간이 넓어질수록 완료 요청 수가 증가한다.",
    )

    doc.add_heading("4.2 semantic vs non-semantic", level=2)
    comparison = []
    for policy in ["LinUCB", "SA", "Greedy", "MCTS"]:
        sem = float(metric_row(summary_rows, "semantic", 300, policy)["finished"])
        non = float(metric_row(summary_rows, "nonsemantic", 300, policy)["finished"])
        gain = (sem - non) / non * 100.0
        comparison.append([policy, f"{sem:.1f}", f"{non:.1f}", f"{gain:.1f}%"])
    add_table(
        doc,
        ["정책", "Semantic 완료 요청", "Non-semantic 완료 요청", "증가율"],
        comparison,
        [1900, 2500, 2800, 2160],
        font_size=8.4,
    )
    add_figure(
        doc,
        FIG_DIR / "semantic_vs_nonsemantic_300m.png",
        "그림 2. 300 x 300 m 환경에서 semantic processing은 모든 비교 정책에서 non-semantic보다 완료 요청 수가 많다.",
    )

    doc.add_heading("4.3 아직 다른 부분", level=2)
    largest = sorted(final_verify, key=lambda row: float(row["relative_error"]), reverse=True)[:6]
    add_table(
        doc,
        ["Figure", "Area", "Policy", "Metric", "Paper est.", "Repro.", "Rel. err."],
        [
            [
                row["check"],
                row["area"],
                row["policy"],
                row["metric"],
                fmt_num(row["paper_visual_estimate"], 1),
                fmt_num(row["reproduction"], 1),
                fmt_num(row["relative_error"], 2),
            ]
            for row in largest
        ],
        [1700, 850, 1150, 1750, 1300, 1300, 1310],
        font_size=7.7,
    )
    add_figure(
        doc,
        FIG_DIR / "average_time_cost.png",
        "그림 3. 평균 처리 시간. 일부 정책, 특히 SA와 LinUCB의 시간 scale은 논문 추정값보다 크게 남아 있다.",
    )
    add_callout(
        doc,
        "결과 해석",
        "방향성은 재현됐다. semantic 방식이 non-semantic보다 낫고, 일부 정책의 면적 증가 경향도 맞는다. "
        "그러나 stochastic 완료 요청 수와 평균 처리 시간은 논문 그래프와 크게 다르다. "
        "따라서 최종 결과는 exact reproduction이 아니라 공개 정보 기반 보정 재현으로 해석해야 한다.",
        fill=CALLOUT_FILL,
    )


def write_limitations(doc: Document) -> None:
    doc.add_heading("5. 한계점", level=1)
    add_bullet(doc, "공식 AirTalking source code와 raw result가 공개되어 있지 않아 그래프의 정확한 수치를 재계산할 수 없다.")
    add_bullet(doc, "request probability, workload distribution, P_move/P_hover, codec power, interference scheduling, policy hyperparameter가 논문에 수치로 공개되어 있지 않다.")
    add_bullet(doc, "Cityscapes는 공개 데이터라 실제로 사용했지만, 논문이 언급한 모델 weight와 feature tensor 정의는 공개되지 않아 proxy를 사용했다.")
    add_bullet(doc, "논문 그래프의 paper-side 값은 원 데이터가 아니라 그림에서 읽은 visual estimate다.")
    add_bullet(doc, "보정값은 관행적/물리적으로 그럴듯한 범위와 결과 scale을 맞추기 위한 값이지, 논문 저자의 실제 hidden setting이라고 주장할 수 없다.")
    add_callout(
        doc,
        "발표/제출 시 권장 표현",
        "이 실험은 Cityscapes 공개 데이터와 논문 공개 파라미터를 사용한 approximate reproduction이다. "
        "공개되지 않은 simulator 설정은 별도 가정값으로 분리했고, 그 가정 때문에 일부 결과는 논문과 정량적으로 다르다.",
        fill=CAUTION_FILL,
    )

    add_para(
        doc,
        "재현성을 위해 확인해야 할 산출물은 다음과 같다. "
        "Cityscapes 측정값은 cityscapes_semantic_summary.json, 최종 시뮬레이션 설정은 run_metadata.json, "
        "논문 대비 검증은 verification_against_paper_calibrated_final_p012.csv에 들어 있다.",
        size=10.2,
    )


def audit_docx_structure(path: Path) -> dict[str, int | list[str]]:
    with ZipFile(path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
    return {
        "tables": document_xml.count("<w:tbl>"),
        "images": document_xml.count("<w:drawing>"),
        "header_rows": document_xml.count("<w:tblHeader"),
        "image_alt_titles": document_xml.count("descr="),
        "required_sections_found": [
            title
            for title in ["1. 문제정의", "2. 제안내용", "3. 실험내용", "4. 실험결과", "5. 한계점"]
            if title in document_xml
        ],
    }


def build_report() -> Path:
    require_inputs()
    city = json.loads(CITY_SUMMARY_JSON.read_text(encoding="utf-8"))
    metadata = json.loads(METADATA_JSON.read_text(encoding="utf-8"))
    summary_rows = load_csv(SUMMARY_CSV)
    original_verify = load_csv(ORIGINAL_VERIFY_CSV)
    final_verify = load_csv(VERIFY_CSV)

    doc = Document()
    configure_doc(doc)
    add_masthead(doc, metadata, city)
    write_problem_definition(doc)
    write_proposal(doc)
    write_experiment(doc, metadata, city)
    write_results(doc, summary_rows, original_verify, final_verify)
    write_limitations(doc)

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_report()
    audit = audit_docx_structure(path)
    print(path)
    print(json.dumps(audit, ensure_ascii=False, indent=2))
