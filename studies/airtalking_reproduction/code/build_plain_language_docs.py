from __future__ import annotations

import csv
import json
from collections import Counter
from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path(__file__).resolve().parents[3]
ADAPT_DIR = WORKSPACE_ROOT / "studies" / "adaptive_semantic_compression"
ADAPT_OUT = ADAPT_DIR / "results" / "probe_outputs"
REPRO_DIR = WORKSPACE_ROOT / "studies" / "airtalking_reproduction" / "results" / "airtalking_cityscapes_calibrated_final_p012"
CITY_DIR = WORKSPACE_ROOT / "studies" / "airtalking_reproduction" / "results" / "cityscapes_semantic_measurement"

PROPOSAL_PATH = ADAPT_DIR / "reports" / "Adaptive_Semantic_Compression_Proposal_EASY_KR.docx"
IMPLEMENTATION_PATH = WORKSPACE_ROOT / "studies" / "airtalking_reproduction" / "reports" / "AirTalking_Experiment_Implementation_EASY_KR.docx"

LATIN_FONT = "Calibri"
KOREAN_FONT = "Malgun Gothic"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "555555"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
POSITIVE_FILL = "EEF7F1"
CAUTION_FILL = "FFF8E8"
BORDER = "D7DBE2"


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = LATIN_FONT
    style._element.get_or_add_rPr()
    rfonts = style._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        style._element.rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_first_row_as_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(table, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, size: float = 8.8, bold: bool = False, color: str = "000000", align=None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.12
    if align is not None:
        para.alignment = align
    run = para.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], font_size: float = 8.6) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    mark_first_row_as_header(table)
    set_cell_margins(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, header, size=font_size, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(values) > 2 else None
            set_cell_text(cells[idx], value, size=font_size, align=align)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    mark_first_row_as_header(table)
    set_cell_margins(table, top=155, bottom=155, start=190, end=190)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    set_run_font(r1, size=10.6, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.8, color="222222")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_para(doc: Document, text: str, size: float = 10.6, color: str = "000000", bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.4)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=8.6, color=MUTED, italic=True)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.05) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image = p.add_run().add_picture(str(path), width=Inches(width))
    image._inline.docPr.set("descr", caption)
    image._inline.docPr.set("title", caption.split(".", 1)[0])
    add_caption(doc, caption)


def configure_doc(doc: Document, header_text: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    set_style_font(styles["Heading 1"], size=16, color=BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Heading 2"], size=13, color=BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    set_style_font(styles["List Bullet"], size=10.4)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Bullet"].paragraph_format.space_after = Pt(8)
    styles["List Bullet"].paragraph_format.line_spacing = 1.167
    hp = section.header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run(header_text)
    set_run_font(hr, size=8.5, color=MUTED)
    fp = section.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Generated: " + date.today().isoformat())
    set_run_font(fr, size=8.5, color=MUTED)


def title_block(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=23, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_run_font(r, size=13.2, color=MUTED, bold=True)


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def verdict_counts(path: Path) -> Counter:
    return Counter(row["verdict"] for row in load_csv(path))


def build_easy_proposal() -> Path:
    policy = json.loads((ADAPT_OUT / "policy_summary.json").read_text(encoding="utf-8"))
    quality = load_csv(ADAPT_OUT / "compression_quality.csv")
    doc = Document()
    configure_doc(doc, "쉬운 연구 제안서 | Adaptive semantic compression")
    title_block(
        doc,
        "연구 제안서 - 쉬운 설명판",
        "통신 상태에 맞춰 사진의 의미 정보를 얼마나 줄일지 자동으로 고르는 연구",
    )
    add_callout(
        doc,
        "교수님께 드리는 한 줄 요청",
        "AirTalking 논문의 한계인 '항상 같은 비율로 의미 정보를 줄이는 방식'을 개선해서, 통신 상태에 따라 압축 강도를 바꾸는 연구 주제로 진행해도 될지 검토받고자 합니다.",
        fill=POSITIVE_FILL,
    )

    doc.add_heading("1. 이 연구를 왜 하려는가", level=1)
    add_para(doc, "기존 AirTalking 논문은 드론끼리 데이터를 주고받을 때 사진 전체를 보내지 말고, 사진에서 중요한 내용만 뽑아 보내면 더 효율적이라고 봅니다.")
    add_para(doc, "그런데 기존 논문은 의미 정보를 줄이는 비율을 거의 고정값처럼 사용합니다. 통신이 잘 될 때와 안 될 때를 나누어 압축 강도를 바꾸지는 않습니다.")
    add_callout(
        doc,
        "쉽게 비유하면",
        "영상 통화에서 인터넷이 느리면 화질을 낮추고, 인터넷이 빠르면 화질을 높이는 것과 비슷합니다. 이 연구는 드론 통신에서도 그런 식으로 '보낼 정보의 양'을 상황에 맞게 조절해보자는 아이디어입니다.",
    )

    doc.add_heading("2. 어려운 용어 먼저 풀어쓰기", level=1)
    add_table(
        doc,
        ["용어", "쉬운 뜻", "이 연구에서 하는 역할"],
        [
            ["Semantic", "사진의 의미, 즉 도로/차/사람 같은 중요한 내용", "사진 전체 대신 중요한 내용만 보내기 위해 사용"],
            ["Compression", "데이터 크기를 줄이는 것", "작게 만들수록 빨리 보낼 수 있음"],
            ["Adaptive", "상황에 맞춰 바꾸는 것", "통신 상태에 따라 압축 강도를 바꿈"],
            ["Encoder", "사진에서 의미 정보를 뽑는 장치 또는 모델", "큰 사진을 작은 의미 정보로 바꿈"],
            ["Decoder", "받은 의미 정보를 다시 쓸 수 있게 만드는 장치 또는 모델", "전송 후 의미 정보를 복원하거나 해석함"],
            ["SINR", "통신 상태 점수", "신호가 좋으면 큰 값, 나쁘면 작은 값"],
            ["mIoU", "의미 정보가 얼마나 정확히 남았는지 보는 점수", "압축을 세게 해도 의미가 보존되는지 확인"],
        ],
        [1700, 3800, 3860],
        font_size=7.9,
    )

    doc.add_heading("3. 제안하는 연구 주제", level=1)
    add_para(doc, "주제명은 '채널 상태 기반 적응형 의미 정보 압축'으로 잡을 수 있습니다. 영어로는 Channel-aware Adaptive Semantic Compression입니다.")
    add_table(
        doc,
        ["통신 상태", "보낼 정보량", "이유"],
        [
            ["나쁨", "아주 작게 보냄", "전송이 실패하거나 너무 오래 걸리는 것을 줄이기 위해"],
            ["보통", "중간 정도로 줄임", "시간과 품질을 균형 있게 맞추기 위해"],
            ["좋음", "덜 줄이고 품질을 보존", "통신 여유가 있으므로 더 정확한 의미 정보를 보내기 위해"],
        ],
        [1800, 2600, 4960],
        font_size=8.4,
    )

    doc.add_heading("4. 어떤 기술을 쓰는가", level=1)
    add_table(
        doc,
        ["기술", "쉽게 말하면", "어디에 쓸 것인가"],
        [
            ["컴퓨터비전", "컴퓨터가 사진을 이해하게 하는 분야", "Cityscapes 도로 사진과 의미 라벨 처리"],
            ["딥러닝", "사진에서 패턴을 배워 의미를 찾는 AI", "추후 U-Net, SegFormer 같은 공개 모델을 encoder 후보로 사용"],
            ["무선통신 계산", "신호가 좋은지 나쁜지 계산", "드론 사이 전송 속도와 지연시간 계산"],
            ["시뮬레이션", "실제 드론 대신 코드 안에서 실험", "여러 드론, 여러 기기, 여러 통신 상태를 반복 실험"],
            ["스케줄링", "누가 누구에게 어떤 방식으로 보낼지 정하는 과정", "압축 강도와 드론 선택을 함께 결정"],
        ],
        [1700, 3600, 4060],
        font_size=8.0,
    )

    doc.add_heading("5. 실험을 어떻게 할 것인가", level=1)
    add_table(
        doc,
        ["순서", "할 일", "나오는 결과"],
        [
            ["1", "Cityscapes 사진과 정답 라벨을 준비", "도로/차/사람 위치 정보"],
            ["2", "의미 정보를 여러 크기로 줄여봄", "강한 압축, 중간 압축, 약한 압축"],
            ["3", "줄인 뒤에도 의미가 잘 남았는지 확인", "mIoU, pixel accuracy"],
            ["4", "드론 통신 시뮬레이터에 넣음", "완료 요청 수, 평균 시간, 에너지"],
            ["5", "기존 방식과 제안 방식을 비교", "고정 압축 vs 적응형 압축"],
        ],
        [850, 5050, 3460],
        font_size=8.1,
    )

    doc.add_heading("6. 이미 해본 예비 실험", level=1)
    add_para(doc, "Cityscapes 의미 라벨 400개를 사용해, 의미 정보를 여러 크기로 줄였을 때 데이터 크기와 의미 품질이 어떻게 바뀌는지 확인했습니다.")
    rows = []
    for row in quality:
        rows.append(
            [
                row["mode"],
                f"{float(row['feature_ratio_mean']) * 100:.2f}%",
                f"{float(row['mean_iou_mean']):.3f}",
                f"{float(row['pixel_accuracy_mean']):.3f}",
            ]
        )
    add_table(doc, ["압축 모드", "원본 대비 크기", "의미 품질", "픽셀 정확도"], rows, [1800, 2100, 1700, 3760], font_size=8.4)
    add_figure(doc, ADAPT_OUT / "figures" / "quality_vs_payload.png", "그림 1. 데이터 크기를 줄이면 전송은 쉬워지지만 의미 품질은 조금씩 낮아진다.")
    ps = policy["policy_summary"]
    add_table(
        doc,
        ["방식", "평균 전송시간", "쉬운 해석"],
        [
            ["원본 전송", f"{ps['raw_time_s']['mean']:.2f}초", "사진 전체를 보내서 가장 오래 걸림"],
            ["기존 논문식 고정 압축", f"{ps['fixed_paper_time_s']['mean']:.2f}초", "항상 비슷한 크기로 줄임"],
            ["제안 방식", f"{ps['adaptive_time_s']['mean']:.2f}초", "통신 상태에 따라 줄이는 정도를 바꿈"],
        ],
        [2200, 2000, 5160],
        font_size=8.4,
    )
    add_callout(
        doc,
        "중요한 주의",
        f"제안 방식이 고정 압축보다 평균 전송시간을 약 {ps['adaptive_vs_fixed_paper_time_reduction_pct']:.1f}% 줄이는 것으로 나왔지만, 이것은 예비 계산입니다. 실제 논문에서는 드론 시뮬레이터 전체에 넣어 다시 검증해야 합니다.",
        fill=CAUTION_FILL,
    )
    add_figure(doc, ADAPT_OUT / "figures" / "delivery_time_by_policy.png", "그림 2. 예비 계산에서는 상황에 맞춰 압축하는 방식이 더 빠르게 나왔다.")

    doc.add_heading("7. 교수님께 확인받고 싶은 점", level=1)
    add_bullet(doc, "이 주제가 학부생 학회 제출용으로 적절한지")
    add_bullet(doc, "처음에는 Cityscapes label 기반 실험으로 진행하고, 시간이 되면 공개 딥러닝 모델까지 확장해도 되는지")
    add_bullet(doc, "기여점을 '적응형 압축 정책'에 집중하는 것이 좋은지, 아니면 encoder 구현까지 포함해야 하는지")
    add_bullet(doc, "비교 실험에서 꼭 넣어야 할 baseline과 평가 지표가 무엇인지")
    doc.save(PROPOSAL_PATH)
    return PROPOSAL_PATH


def build_easy_implementation() -> Path:
    city = json.loads((CITY_DIR / "cityscapes_semantic_summary.json").read_text(encoding="utf-8"))
    meta = json.loads((REPRO_DIR / "run_metadata.json").read_text(encoding="utf-8"))
    counts = verdict_counts(REPRO_DIR / "verification_against_paper_calibrated_final_p012.csv")
    doc = Document()
    configure_doc(doc, "쉬운 실험 구현서 | AirTalking reproduction")
    title_block(doc, "AirTalking 재현 실험 구현서 - 쉬운 설명판", "무슨 데이터를 넣고, 코드가 무엇을 계산했고, 결과를 어떻게 봤는지 설명")
    add_callout(
        doc,
        "한 줄 요약",
        "Cityscapes 도로 사진에서 '의미 정보가 원본보다 훨씬 작다'는 것을 계산한 뒤, 그 값을 드론 통신 가상 실험에 넣어 원본 전송과 의미 정보 전송을 비교했습니다.",
        fill=POSITIVE_FILL,
    )

    doc.add_heading("1. 이 실험의 목적", level=1)
    add_para(doc, "논문은 드론이 사진 전체를 보내는 대신, 사진에서 중요한 내용만 뽑아 보내면 더 빠르고 효율적이라고 주장합니다.")
    add_para(doc, "우리는 이 주장이 맞는지 보기 위해 실제 드론을 날린 것이 아니라, 컴퓨터 코드 안에 드론 환경을 만들고 반복 실험했습니다.")

    doc.add_heading("2. 사용한 데이터", level=1)
    add_table(
        doc,
        ["데이터", "쉬운 설명", "어디에 사용했나"],
        [
            ["Cityscapes leftImg8bit", "도시 도로 원본 사진", "원본 사진 크기 계산"],
            ["Cityscapes gtFine labelIds", "사진에서 도로/차/사람 위치를 표시한 정답 지도", "의미 정보 크기 계산"],
            ["train/val 3,475쌍", "학습/검증용 공개 데이터", "실제 계산에 사용"],
            ["test split", "정답 라벨이 제대로 없어 사용하지 않음", "실험 입력에서 제외"],
        ],
        [2300, 3900, 3160],
        font_size=8.3,
    )

    doc.add_heading("3. 첫 번째 계산: 사진 전체 vs 의미 정보", level=1)
    raw_mb = city["raw_uncompressed_bytes_mean"] / 1024 / 1024
    sem_mb = city["semantic_feature_bytes_mean"] / 1024 / 1024
    add_table(
        doc,
        ["항목", "값", "뜻"],
        [
            ["원본 사진 크기", f"{raw_mb:.2f} MB", "색깔 정보가 모두 들어 있는 사진"],
            ["의미 정보 크기", f"{sem_mb:.2f} MB", "도로/차/사람 같은 의미만 남긴 정보"],
            ["원본 대비 비율", f"{city['rho_c_feature_uncompressed_mean']:.6f}", "원본의 약 10.4% 크기"],
            ["논문값", "0.104", "우리 측정값과 거의 같음"],
        ],
        [2200, 1900, 5260],
        font_size=8.5,
    )
    add_callout(doc, "쉽게 해석", "사진 전체를 100장짜리 책이라고 보면, 의미 정보는 약 10장짜리 요약본에 가깝습니다. 그래서 통신으로 보낼 양이 크게 줄어듭니다.")

    doc.add_heading("4. 두 번째 계산: 드론 가상 실험장 만들기", level=1)
    paper = meta["paper_params"]
    add_table(
        doc,
        ["설정", "값", "쉬운 설명"],
        [
            ["드론 수", str(paper["n_uav"]), "데이터를 중계하는 공중 장치"],
            ["지상 기기 수", str(paper["n_device"]), "데이터를 보내거나 받는 사용자 기기"],
            ["실험 시간", f"{paper['t_slots']}초", "코드 안에서 1초 단위로 진행"],
            ["드론 높이", "평균 20m", "지상 기기는 2D 평면, 드론은 높이를 붙여 3D 거리 계산"],
            ["공간 크기", "100m x 100m부터 500m x 500m", "좁은 곳과 넓은 곳을 비교"],
        ],
        [2200, 2000, 5160],
        font_size=8.4,
    )
    add_para(doc, "코드 안에서는 지상 기기가 요청을 만듭니다. 예를 들면 '기기 A의 데이터를 기기 B에게 보내줘' 같은 요청입니다. 드론은 그 요청을 받아 중간에서 데이터를 전달합니다.")

    doc.add_heading("5. 비교한 두 가지 전송 방식", level=1)
    add_table(
        doc,
        ["방식", "무엇을 보내나", "장점", "단점"],
        [
            ["원본 전송", "사진 전체", "정보가 그대로 있음", "데이터가 커서 오래 걸림"],
            ["의미 정보 전송", "도로/차/사람 같은 핵심 정보", "작아서 빨리 보낼 수 있음", "encoder/decoder 처리 시간이 추가됨"],
        ],
        [1700, 3000, 2400, 2260],
        font_size=8.2,
    )
    add_callout(doc, "encoder와 decoder", "Encoder는 사진에서 중요한 내용을 뽑는 역할입니다. Decoder는 받은 의미 정보를 다시 쓸 수 있게 만드는 역할입니다. 우리는 모델을 새로 학습하지 않고, 논문에 공개된 처리 속도를 시뮬레이션에 넣었습니다.", fill=CAUTION_FILL)

    doc.add_heading("6. 드론 통신 속도는 어떻게 계산했나", level=1)
    add_bullet(doc, "드론끼리 가까우면 신호가 좋고 전송이 빠릅니다.")
    add_bullet(doc, "드론끼리 멀거나 동시에 여러 통신이 있으면 신호가 나빠지고 전송이 느립니다.")
    add_bullet(doc, "코드는 거리, 신호 세기, 간섭을 계산해서 전송 속도를 구합니다.")
    add_bullet(doc, "실제 드론 비행 실험이 아니라, 논문 수식을 바탕으로 한 가상 실험입니다.")

    doc.add_heading("7. 정책 알고리즘은 무엇인가", level=1)
    add_para(doc, "요청이 들어오면 어떤 드론이 보내고 받을지 정해야 합니다. 이 선택 방법을 정책 알고리즘이라고 부릅니다.")
    add_table(
        doc,
        ["이름", "쉬운 설명"],
        [
            ["Stochastic", "거의 랜덤으로 고르는 기준점"],
            ["LinUCB", "이전 결과를 보고 좋아 보이는 선택을 더 자주 고르는 방식"],
            ["SA", "여러 후보를 바꿔보며 괜찮은 답을 찾는 방식"],
            ["Greedy", "지금 당장 가장 좋아 보이는 선택을 고르는 방식"],
            ["MCTS", "앞으로 어떻게 될지 여러 번 가정해보고 고르는 방식"],
        ],
        [2200, 7160],
        font_size=8.4,
    )

    doc.add_heading("8. 결과는 논문과 얼마나 맞았나", level=1)
    add_table(
        doc,
        ["판정", "개수", "뜻"],
        [
            ["Match", str(counts.get("match", 0)), "논문 그래프와 꽤 가까움"],
            ["Partial", str(counts.get("partial", 0)), "방향은 비슷하지만 차이가 있음"],
            ["Mismatch", str(counts.get("mismatch", 0)), "숫자가 많이 다름"],
        ],
        [1900, 1500, 5960],
        font_size=8.6,
    )
    add_para(doc, "결론적으로 완전 일치는 아닙니다. 하지만 의미 정보를 보내는 방식이 원본 전송보다 유리하다는 방향성은 재현됐습니다.")
    add_figure(doc, REPRO_DIR / "figures" / "semantic_vs_nonsemantic_300m.png", "그림 1. 300m 환경에서 의미 정보 전송이 원본 전송보다 더 많은 요청을 처리하는 경향.")

    doc.add_heading("9. 왜 완전히 같지는 않은가", level=1)
    add_bullet(doc, "논문이 encoder/decoder 코드와 학습된 모델 파일을 공개하지 않았습니다.")
    add_bullet(doc, "요청이 얼마나 자주 생기는지, 데이터 크기 분포, 드론 전력값 같은 핵심 설정이 공개되지 않았습니다.")
    add_bullet(doc, "논문 그래프의 원본 숫자가 없어 그림을 보고 대략 읽어 비교했습니다.")
    add_bullet(doc, "따라서 이 결과는 정확 복제가 아니라 공개 정보 기반 근사 재현입니다.")

    doc.add_heading("10. 파일 위치", level=1)
    add_table(
        doc,
        ["파일/폴더", "역할"],
        [
            ["studies/airtalking_reproduction/results/cityscapes_semantic_measurement", "Cityscapes로 의미 정보 크기를 측정한 결과"],
            ["studies/airtalking_reproduction/results/airtalking_cityscapes_calibrated_final_p012", "최종 보정 실험 결과"],
            ["studies/airtalking_reproduction/code/airtalking_reproduction.py", "드론 통신 가상 실험 코드"],
            ["studies/airtalking_reproduction/code/measure_cityscapes_semantics.py", "원본 사진과 의미 정보 크기 측정 코드"],
            ["studies/airtalking_reproduction/code/verify_against_paper.py", "논문 그래프와 결과 비교 코드"],
        ],
        [3600, 5760],
        font_size=8.3,
    )
    doc.save(IMPLEMENTATION_PATH)
    return IMPLEMENTATION_PATH


def audit_docx(path: Path) -> dict[str, int]:
    with ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return {
        "tables": xml.count("<w:tbl>"),
        "images": xml.count("<w:drawing>"),
        "headers": xml.count("<w:tblHeader"),
        "alt_text": xml.count("descr="),
    }


def main() -> None:
    proposal = build_easy_proposal()
    implementation = build_easy_implementation()
    print(json.dumps({
        "proposal": str(proposal),
        "proposal_audit": audit_docx(proposal),
        "implementation": str(implementation),
        "implementation_audit": audit_docx(implementation),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
