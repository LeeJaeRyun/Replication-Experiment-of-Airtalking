from __future__ import annotations

import base64
import csv
import hashlib
import importlib.util
import json
import math
import statistics
import sys
import tempfile
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "tools" / "finalize_research_reports.py"
SPEC = importlib.util.spec_from_file_location("finalize_research_reports", MODULE_PATH)
assert SPEC is not None and SPEC.loader is not None
finalizer = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = finalizer
SPEC.loader.exec_module(finalizer)


class FinalizeResearchReportsTests(unittest.TestCase):
    @staticmethod
    def _write_csv(path: Path, fieldnames: list[str], rows: list[dict[str, object]]) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.DictWriter(handle, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(rows)

    @staticmethod
    def _write_json(path: Path, payload: dict[str, object]) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2, allow_nan=False),
            encoding="utf-8",
        )

    @staticmethod
    def _metric_base(mode: str, area: int, policy_index: int, metric_index: int) -> float:
        offsets = {
            "semantic": 20.0,
            "nonsemantic": 10.0,
            "fixed_paper_like": 30.0,
            "adaptive_semantic": 31.0,
        }
        value = offsets[mode] + area / 1000.0 + policy_index / 100.0 + metric_index / 10000.0
        if finalizer.CORE_REPEAT_METRICS[metric_index] == "avg_time":
            value = 5.0 if mode in {"semantic", "fixed_paper_like"} else 4.0
        if finalizer.CORE_REPEAT_METRICS[metric_index] == "semantic_quality":
            value = 0.90 if mode in {"semantic", "fixed_paper_like"} else 0.89
        return value

    @classmethod
    def _experiment_tables(
        cls, keys: set[tuple[str, int, str]], repeats: int = 2
    ) -> tuple[list[dict[str, object]], list[dict[str, object]], list[dict[str, object]]]:
        summary_rows: list[dict[str, object]] = []
        repeat_rows: list[dict[str, object]] = []
        statistics_rows: list[dict[str, object]] = []
        for mode, area, policy in sorted(keys):
            policy_index = finalizer.POLICIES.index(policy)
            summary: dict[str, object] = {"mode": mode, "area": area, "policy": policy}
            group_repeats: list[dict[str, object]] = []
            for repeat in range(repeats):
                row: dict[str, object] = {
                    "mode": mode,
                    "area": area,
                    "policy": policy,
                    "repeat": repeat,
                }
                for metric_index, metric in enumerate(finalizer.CORE_REPEAT_METRICS):
                    center = cls._metric_base(mode, area, policy_index, metric_index)
                    row[metric] = center + (-0.5 if repeat == 0 else 0.5)
                group_repeats.append(row)
                repeat_rows.append(row)

            statistics_row: dict[str, object] = {
                "mode": mode,
                "area": area,
                "policy": policy,
                "repeat_count": repeats,
                "ci95_method": finalizer.CI95_METHOD,
            }
            for metric in finalizer.CORE_REPEAT_METRICS:
                values = [float(row[metric]) for row in group_repeats]
                mean = statistics.fmean(values)
                std = statistics.stdev(values) if repeats > 1 else 0.0
                margin = finalizer._student_t_critical_95(repeats) * std / math.sqrt(repeats)
                summary[metric] = mean
                statistics_row.update(
                    {
                        f"{metric}_n": repeats,
                        f"{metric}_mean": mean,
                        f"{metric}_std": std,
                        f"{metric}_ci95_margin": margin,
                        f"{metric}_ci95_low": mean - margin,
                        f"{metric}_ci95_high": mean + margin,
                    }
                )
            summary_rows.append(summary)
            statistics_rows.append(statistics_row)
        return summary_rows, repeat_rows, statistics_rows

    @classmethod
    def _write_experiment_tables(
        cls,
        directory: Path,
        keys: set[tuple[str, int, str]],
        repeats: int = 2,
    ) -> tuple[list[dict[str, object]], list[dict[str, object]], list[dict[str, object]]]:
        summary, repeat_rows, statistics_rows = cls._experiment_tables(keys, repeats)
        summary_fields = ["mode", "area", "policy", *finalizer.CORE_REPEAT_METRICS]
        repeat_fields = ["mode", "area", "policy", "repeat", *finalizer.CORE_REPEAT_METRICS]
        statistics_fields = ["mode", "area", "policy", "repeat_count", "ci95_method"]
        statistics_fields.extend(
            f"{metric}_{suffix}"
            for metric in finalizer.CORE_REPEAT_METRICS
            for suffix in ("n", "mean", "std", "ci95_margin", "ci95_low", "ci95_high")
        )
        cls._write_csv(directory / "summary_metrics.csv", summary_fields, summary)
        cls._write_csv(directory / "repeat_metrics.csv", repeat_fields, repeat_rows)
        cls._write_csv(directory / "statistical_summary.csv", statistics_fields, statistics_rows)
        return summary, repeat_rows, statistics_rows

    @classmethod
    def _make_strict_fixture(cls, root: Path) -> finalizer.InputPaths:
        enhanced = root / "enhanced"
        reproduction = root / "reproduction"
        adaptive = root / "adaptive"
        baseline_neural = root / "baseline_neural"
        baseline_reproduction = root / "baseline_reproduction"
        baseline_adaptive = root / "baseline_adaptive"
        for directory in (
            enhanced,
            reproduction,
            adaptive,
            baseline_neural,
            baseline_reproduction,
            baseline_adaptive,
        ):
            directory.mkdir(parents=True)

        for filename in (
            "best_checkpoint.pt",
            "last_checkpoint.pt",
            "final_checkpoint.pt",
            "launch_manifest.json",
            "training_source.py",
        ):
            (enhanced / filename).write_bytes(f"synthetic {filename}".encode())
        cls._write_json(
            enhanced / "result_summary.json",
            {
                "status": "completed",
                "dataset": {
                    "name": "synthetic Cityscapes-like",
                    "train_samples": 4,
                    "val_samples": 2,
                    "image_size": {"width": 256, "height": 128},
                },
                "model": {
                    "encoder": "SyntheticEncoder",
                    "decoder": "SyntheticDecoder",
                    "paper_like_active_channels": 20,
                },
                "training": {"epochs_requested": 2, "best_epoch": 2},
                "paper_like_profile": {
                    "mean_iou": 0.5,
                    "pixel_accuracy": 0.8,
                    "psnr_db": 22.0,
                    "ssim": 0.7,
                },
                "timing": {
                    "encode_including_8bit_fake_quantization": {"median_ms": 2.0},
                    "decode_from_latent_only": {"median_ms": 3.0},
                },
                "environment": {"device": "cpu", "command_windows": "python synthetic.py"},
                "elapsed_seconds": 10.0,
                "artifacts": {
                    "best_checkpoint": "best_checkpoint.pt",
                    "last_resume_checkpoint": "last_checkpoint.pt",
                    "final_training_checkpoint": "final_checkpoint.pt",
                },
                "provenance": {
                    "launch_manifest": "launch_manifest.json",
                    "training_source_snapshot": "training_source.py",
                },
            },
        )
        rates = []
        for channel in (20, 40, 60, 80, 120):
            rates.append(
                {
                    "active_channels": channel,
                    "measured_rho_uint8_over_raw_rgb": channel / 1000,
                    "measured_rho_zlib_over_raw_rgb": channel / 2000,
                    "mean_iou": 0.5,
                    "pixel_accuracy": 0.8,
                    "psnr_db": 22.0,
                    "ssim": 0.7,
                }
            )
        cls._write_csv(enhanced / "rate_quality.csv", list(rates[0]), rates)
        cls._write_csv(
            enhanced / "training_history.csv",
            ["epoch", "train_loss", "val_loss", "val_mean_iou", "val_pixel_accuracy"],
            [
                {"epoch": 1, "train_loss": 1.0, "val_loss": 0.9, "val_mean_iou": 0.4, "val_pixel_accuracy": 0.7},
                {"epoch": 2, "train_loss": 0.8, "val_loss": 0.7, "val_mean_iou": 0.5, "val_pixel_accuracy": 0.8},
            ],
        )
        confusion_fields = ["label", *(f"c{index}" for index in range(19))]
        confusion_rows = []
        for row_index in range(19):
            row: dict[str, object] = {"label": f"class-{row_index}"}
            row.update({f"c{column}": 1 if column == row_index else 0 for column in range(19)})
            confusion_rows.append(row)
        cls._write_csv(
            enhanced / "confusion_matrix_paper_like.csv",
            confusion_fields,
            confusion_rows,
        )
        enhanced_profile = {
            "schema_version": 2,
            "paper_like_active_channels": 20,
            "rho_c_feature_uncompressed_mean": 0.02,
            "rho_c_feature_zlib_mean": 0.01,
            "semantic_quality_miou_final": 0.5,
            "pixel_accuracy_final": 0.8,
            "multi_rate_profiles": [{"active_channels": channel} for channel in (20, 40, 60, 80, 120)],
        }
        cls._write_json(enhanced / "airtalking_semantic_summary.json", enhanced_profile)

        cls._write_experiment_tables(reproduction, finalizer._expected_reproduction_keys())
        cls._write_csv(
            reproduction / "verification_against_paper_latest.csv",
            ["check", "metric", "paper_visual_estimate", "reproduction", "verdict"],
            [{"check": "figure", "metric": "finished", "paper_visual_estimate": 1.0, "reproduction": 1.0, "verdict": "match"}],
        )
        reproduction_summary = reproduction / "summary_metrics.csv"
        cls._write_json(
            reproduction / "verification_against_paper_latest.json",
            {
                "schema_version": 2,
                "status": "completed",
                "source_summary": {
                    "path": "summary_metrics.csv",
                    "sha256": hashlib.sha256(reproduction_summary.read_bytes()).hexdigest(),
                },
                "row_count": 1,
                "verdict_counts": {"match": 1, "partial": 0, "mismatch": 0},
                "qualitative_row_count": 1,
                "qualitative_verdict_counts": {"match": 1, "partial": 0, "mismatch": 0},
            },
        )
        cls._write_json(
            reproduction / "run_metadata.json",
            {
                "paper_params": {"repeats": 2, "t_slots": 10},
                "assumed_params": {"seed": 1, "request_probability": 0.1},
                "semantic_profile": {
                    "source": str((enhanced / "airtalking_semantic_summary.json").resolve()),
                    "applied": True,
                },
                "summary_metrics_csv": "summary_metrics.csv",
                "repeat_metrics_csv": "repeat_metrics.csv",
                "statistical_summary_csv": "statistical_summary.csv",
                "elapsed_seconds": 1.0,
            },
        )

        adaptive_summary_rows, _, _ = cls._write_experiment_tables(
            adaptive, finalizer._expected_adaptive_keys()
        )
        usage_rows = [
            {
                "mode": row["mode"],
                "area": row["area"],
                "policy": row["policy"],
                "mode_low_count": 2.0,
                "mode_high_count": 3.0,
            }
            for row in adaptive_summary_rows
        ]
        cls._write_csv(
            adaptive / "compression_mode_usage.csv",
            ["mode", "area", "policy", "mode_low_count", "mode_high_count"],
            usage_rows,
        )
        (adaptive / "source_quality.csv").write_text("quality\n1\n", encoding="utf-8")
        adaptive_summary = adaptive / "summary_metrics.csv"
        adaptive_index = {
            (str(row["mode"]), int(row["area"]), str(row["policy"])): row
            for row in adaptive_summary_rows
        }
        comparisons = []
        for area in finalizer.AREAS:
            for policy in finalizer.POLICIES:
                fixed = adaptive_index[("fixed_paper_like", area, policy)]
                current = adaptive_index[("adaptive_semantic", area, policy)]
                finished_delta = (
                    (float(current["finished"]) - float(fixed["finished"]))
                    / float(fixed["finished"])
                    * 100.0
                )
                time_delta = (
                    (float(current["avg_time"]) - float(fixed["avg_time"]))
                    / float(fixed["avg_time"])
                    * 100.0
                )
                quality_delta = float(current["semantic_quality"]) - float(fixed["semantic_quality"])
                comparisons.append(
                    {
                        "area": area,
                        "policy": policy,
                        "finished_delta_pct": finished_delta,
                        "avg_time_delta_pct": time_delta,
                        "quality_delta": quality_delta,
                        "pass": finished_delta >= 0.0 and time_delta <= 0.0 and quality_delta >= -0.10,
                    }
                )
        cls._write_json(
            adaptive / "result_validation.json",
            {
                "schema_version": 2,
                "source_summary": {
                    "path": "summary_metrics.csv",
                    "sha256": hashlib.sha256(adaptive_summary.read_bytes()).hexdigest(),
                },
                "passed": True,
                "row_count": 75,
                "expected_row_count": 75,
                "expected_combinations": {
                    "modes": list(finalizer.ADAPTIVE_MODES),
                    "areas": list(finalizer.AREAS),
                    "policies": list(finalizer.POLICIES),
                    "count": 75,
                },
                "errors": [],
                "missing_combinations": [],
                "extra_combinations": [],
                "duplicates": [],
                "duplicate_combinations": [],
                "missing_values": [],
                "invalid_canonical_keys": [],
                "invalid_numeric_values": [],
                "nonfinite_values": [],
                "zero_denominators": [],
                "adaptive_vs_fixed": {
                    "evaluated": True,
                    "max_quality_drop": min(item["quality_delta"] for item in comparisons),
                    "comparisons": comparisons,
                },
            },
        )
        cls._write_json(
            adaptive / "run_metadata.json",
            {
                "source_metadata": str((reproduction / "run_metadata.json").resolve()),
                "source_quality": "source_quality.csv",
                "source_neural_encoder_decoder": str((enhanced / "result_summary.json").resolve()),
                "summary_metrics_csv": "summary_metrics.csv",
                "repeat_metrics_csv": "repeat_metrics.csv",
                "statistical_summary_csv": "statistical_summary.csv",
                "compression_mode_usage_csv": "compression_mode_usage.csv",
                "artifacts": {
                    "summary_metrics_csv": {"path": "summary_metrics.csv"},
                    "repeat_metrics_csv": {"path": "repeat_metrics.csv"},
                    "statistical_summary_csv": {"path": "statistical_summary.csv"},
                    "compression_mode_usage_csv": {"path": "compression_mode_usage.csv"},
                },
                "base_paper_params": {"repeats": 2, "t_slots": 10},
                "profiles": {"fixed_paper_like": {}, "adaptive_semantic": {}},
                "elapsed_seconds": 1.0,
                "neural_quality_mode": "selection",
                "neural_encoder_decoder_anchor": {
                    "multi_rate_profiles": [{"active_channels": channel} for channel in (20, 40, 60, 80, 120)]
                },
            },
        )

        cls._write_json(baseline_neural / "result_summary.json", {"status": "completed"})
        cls._write_json(baseline_neural / "airtalking_semantic_summary.json", enhanced_profile)
        baseline_summary = [
            {"mode": "semantic", "area": 100, "policy": "Greedy", "finished": 1.0, "avg_time": 1.0, "flight_energy_per_req": 1.0}
        ]
        cls._write_csv(
            baseline_reproduction / "summary_metrics.csv",
            ["mode", "area", "policy", "finished", "avg_time", "flight_energy_per_req"],
            baseline_summary,
        )
        cls._write_csv(
            baseline_reproduction / "verification_against_paper_baseline.csv",
            ["check", "verdict"],
            [{"check": "baseline", "verdict": "match"}],
        )
        cls._write_csv(
            baseline_adaptive / "summary_metrics.csv",
            ["mode", "area", "policy", "finished", "avg_time"],
            [{"mode": "adaptive_semantic", "area": 100, "policy": "Greedy", "finished": 1.0, "avg_time": 1.0}],
        )
        return finalizer.InputPaths(
            enhanced_dir=enhanced,
            reproduction_dir=reproduction,
            adaptive_dir=adaptive,
            reports_dir=ROOT / "reports",
            output_dir=root / "final",
            baseline_neural_dir=baseline_neural,
            baseline_reproduction_dir=baseline_reproduction,
            baseline_adaptive_dir=baseline_adaptive,
        )

    def test_complete_synthetic_strict_fixture_finalizes_and_embeds_reproduction_statistics(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            paths = self._make_strict_fixture(Path(temporary))
            manifest = finalizer.finalize_reports(paths)

            self.assertEqual(manifest["mode"], "strict")
            self.assertEqual(len(manifest["markdown"]), 3)
            self.assertEqual(len(manifest["docx"]), 3)
            for path in map(Path, manifest["docx"]):
                self.assertTrue(path.is_file())
            reproduction_markdown = next(
                Path(path)
                for path in manifest["markdown"]
                if Path(path).name.startswith("02_")
            )
            rendered = reproduction_markdown.read_text(encoding="utf-8")
            self.assertIn("repeat 원시값에서 독립 검증한 재현 통계", rendered)
            self.assertIn("sample std", rendered)

    def test_json_loader_rejects_nan_and_infinities_immediately(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            for constant in ("NaN", "Infinity", "-Infinity"):
                path = root / f"{constant.replace('-', 'minus')}.json"
                path.write_text(f'{{"value": {constant}}}', encoding="utf-8")
                with self.subTest(constant=constant), self.assertRaisesRegex(
                    finalizer.ReportFinalizationError, "NaN/Infinity"
                ):
                    finalizer._load_json(
                        path,
                        finalizer.Diagnostics(allow_incomplete=False),
                        "비유한 JSON",
                    )

    def test_required_provenance_rejects_blank_and_directory_and_resolves_windows_relative_file(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            directory = root / "한글 폴더"
            directory.mkdir()
            file_path = directory / "결과 파일.csv"
            file_path.write_text("ok", encoding="utf-8")
            self.assertEqual(
                finalizer._resolve_reference(
                    r"한글 폴더\결과 파일.csv", root, require_file=True
                ),
                file_path.resolve(),
            )

            diagnostics = finalizer.Diagnostics(allow_incomplete=False)
            finalizer._check_path_references(
                {"blank": "   ", "directory": str(directory)},
                ("blank", "directory"),
                root,
                diagnostics,
                "provenance",
            )
            with self.assertRaisesRegex(
                finalizer.ReportFinalizationError, "빈 문자열|디렉터리"
            ):
                diagnostics.fail_if_errors()

    def test_strict_fixture_rejects_duplicate_repeat_key(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            paths = self._make_strict_fixture(Path(temporary))
            repeat_path = paths.adaptive_dir / "repeat_metrics.csv"
            with repeat_path.open(newline="", encoding="utf-8") as handle:
                rows = list(csv.DictReader(handle))
                fields = list(rows[0])
            self._write_csv(repeat_path, fields, [*rows, dict(rows[0])])
            with self.assertRaisesRegex(finalizer.ReportFinalizationError, "중복 key"):
                finalizer.load_evidence(paths, finalizer.Diagnostics(False))

    def test_strict_fixture_rejects_missing_repeat(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            paths = self._make_strict_fixture(Path(temporary))
            repeat_path = paths.reproduction_dir / "repeat_metrics.csv"
            with repeat_path.open(newline="", encoding="utf-8") as handle:
                rows = list(csv.DictReader(handle))
                fields = list(rows[0])
            self._write_csv(repeat_path, fields, rows[1:])
            with self.assertRaisesRegex(finalizer.ReportFinalizationError, "repeat 번호"):
                finalizer.load_evidence(paths, finalizer.Diagnostics(False))

    def test_strict_fixture_rejects_statistics_mismatch(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            paths = self._make_strict_fixture(Path(temporary))
            statistics_path = paths.adaptive_dir / "statistical_summary.csv"
            with statistics_path.open(newline="", encoding="utf-8") as handle:
                rows = list(csv.DictReader(handle))
                fields = list(rows[0])
            rows[0]["finished_std"] = str(float(rows[0]["finished_std"]) + 1.0)
            self._write_csv(statistics_path, fields, rows)
            with self.assertRaisesRegex(finalizer.ReportFinalizationError, "독립 재계산 불일치"):
                finalizer.load_evidence(paths, finalizer.Diagnostics(False))

    def test_strict_fixture_rejects_source_summary_hash_mismatch(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            paths = self._make_strict_fixture(Path(temporary))
            validation_path = paths.adaptive_dir / "result_validation.json"
            validation = json.loads(validation_path.read_text(encoding="utf-8"))
            validation["source_summary"]["sha256"] = "0" * 64
            self._write_json(validation_path, validation)
            with self.assertRaisesRegex(finalizer.ReportFinalizationError, "sha256 불일치"):
                finalizer.load_evidence(paths, finalizer.Diagnostics(False))

    def test_reproduction_verifier_requires_one_csv_and_matching_summary_hash(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            paths = self._make_strict_fixture(root)
            manifest_path = paths.reproduction_dir / "verification_against_paper_latest.json"
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest["source_summary"]["sha256"] = "f" * 64
            self._write_json(manifest_path, manifest)
            with self.assertRaisesRegex(finalizer.ReportFinalizationError, "sha256 불일치"):
                finalizer.load_evidence(paths, finalizer.Diagnostics(False))

            paths = self._make_strict_fixture(root / "second")
            (paths.reproduction_dir / "verification_against_paper_old.csv").write_text(
                "check,metric,paper_visual_estimate,reproduction,verdict\nold,x,1,1,match\n",
                encoding="utf-8",
            )
            with self.assertRaisesRegex(finalizer.ReportFinalizationError, "여러 개"):
                finalizer.load_evidence(paths, finalizer.Diagnostics(False))

    def test_adaptive_validator_requires_passed_true(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            paths = self._make_strict_fixture(Path(temporary))
            validation_path = paths.adaptive_dir / "result_validation.json"
            validation = json.loads(validation_path.read_text(encoding="utf-8"))
            validation["passed"] = False
            self._write_json(validation_path, validation)
            with self.assertRaisesRegex(finalizer.ReportFinalizationError, "passed는 반드시 true"):
                finalizer.load_evidence(paths, finalizer.Diagnostics(False))

    def test_usage_contract_rejects_missing_combination_and_blank_count(self) -> None:
        expected = {("adaptive_semantic", 100, "Greedy"), ("adaptive_semantic", 200, "Greedy")}
        diagnostics = finalizer.Diagnostics(allow_incomplete=False)
        finalizer._validate_usage_contract(
            [
                {
                    "mode": "adaptive_semantic",
                    "area": "100",
                    "policy": "Greedy",
                    "mode_low_count": "",
                }
            ],
            expected,
            diagnostics,
        )
        with self.assertRaisesRegex(
            finalizer.ReportFinalizationError, "coverage 불일치|비었거나"
        ):
            diagnostics.fail_if_errors()

    def test_markdown_table_preserves_windows_korean_path_and_only_unescapes_pipe(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            markdown = root / "windows-path.md"
            markdown.write_text(
                "| 경로 | 설명 |\n"
                "|---|---|\n"
                r"| C:\실험 폴더\한글 파일.csv | 값 A \| 값 B |" + "\n"
                + r"| \\서버\공유 폴더\결과.csv | UNC |" + "\n",
                encoding="utf-8",
            )
            docx = root / "windows-path.docx"
            finalizer.markdown_to_docx(markdown, docx, [])

            from docx import Document

            table = Document(docx).tables[0]
            self.assertEqual(table.cell(1, 0).text, r"C:\실험 폴더\한글 파일.csv")
            self.assertEqual(table.cell(1, 1).text, "값 A | 값 B")
            self.assertEqual(table.cell(2, 0).text, r"\\서버\공유 폴더\결과.csv")

    def test_adaptive_claim_uses_minimum_repeat_count_per_combination(self) -> None:
        evidence = finalizer.Evidence(
            adaptive_validation={"passed": True},
            adaptive_metadata={
                "neural_quality_mode": "selection",
                "neural_encoder_decoder_anchor": {"multi_rate_profiles": [{}, {}, {}, {}, {}]},
            },
            adaptive_repeats=[
                {"mode": "adaptive_semantic", "area": "100", "policy": "Greedy", "repeat": "0"},
                {"mode": "adaptive_semantic", "area": "100", "policy": "Greedy", "repeat": "1"},
                {"mode": "adaptive_semantic", "area": "200", "policy": "Greedy", "repeat": "0"},
            ],
        )
        self.assertIn("조합별 최소 repeat 1개", finalizer._render_adaptive_claim(evidence))

    def test_repository_templates_have_exact_33_key_contract(self) -> None:
        templates = finalizer._read_templates(ROOT / "reports")
        keys: list[str] = []
        for _, text in templates.values():
            keys.extend(finalizer.AUTO_RE.findall(text))
        self.assertEqual(len(keys), 33)
        self.assertEqual(set(keys), finalizer.EXPECTED_AUTO_KEYS)

    def test_allow_incomplete_fills_every_auto_and_creates_three_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            enhanced = root / "enhanced"
            reproduction = root / "reproduction"
            adaptive = root / "adaptive"
            output = root / "final"
            enhanced.mkdir()
            reproduction.mkdir()
            adaptive.mkdir()
            paths = finalizer.InputPaths(
                enhanced_dir=enhanced,
                reproduction_dir=reproduction,
                adaptive_dir=adaptive,
                reports_dir=ROOT / "reports",
                output_dir=output,
                baseline_neural_dir=None,
                baseline_reproduction_dir=None,
            )

            first = finalizer.finalize_reports(paths, allow_incomplete=True)
            markdown_before = {
                path.name: path.read_bytes() for path in map(Path, first["markdown"])
            }
            second = finalizer.finalize_reports(paths, allow_incomplete=True)

            self.assertEqual(len(first["markdown"]), 3)
            self.assertEqual(len(first["docx"]), 3)
            self.assertGreater(len(first["warnings"]), 0)
            for raw_path in second["markdown"]:
                path = Path(raw_path)
                text = path.read_text(encoding="utf-8")
                self.assertNotIn("<!-- AUTO:", text)
                self.assertIn("미실행/증거 없음", text)
                self.assertEqual(markdown_before[path.name], path.read_bytes())
            for raw_path in second["docx"]:
                self.assertTrue(Path(raw_path).is_file())

            from docx import Document

            document = Document(second["docx"][0])
            self.assertGreater(len(document.paragraphs), 10)
            self.assertGreater(len(document.tables), 1)
            self.assertIn("인코더", "\n".join(p.text for p in document.paragraphs))

    def test_strict_mode_rejects_missing_core_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            for name in ("enhanced", "reproduction", "adaptive"):
                (root / name).mkdir()
            paths = finalizer.InputPaths(
                enhanced_dir=root / "enhanced",
                reproduction_dir=root / "reproduction",
                adaptive_dir=root / "adaptive",
                reports_dir=ROOT / "reports",
                output_dir=root / "final",
                baseline_neural_dir=None,
                baseline_reproduction_dir=None,
            )
            with self.assertRaisesRegex(
                finalizer.ReportFinalizationError, "result_summary.json"
            ):
                finalizer.finalize_reports(paths)

    def test_malformed_json_is_fatal_even_allow_incomplete(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "bad.json"
            path.write_text('{"status": ', encoding="utf-8")
            diagnostics = finalizer.Diagnostics(allow_incomplete=True)
            with self.assertRaisesRegex(
                finalizer.ReportFinalizationError, "JSON 문법 오류"
            ):
                finalizer._load_json(path, diagnostics, "손상 테스트")

    def test_malformed_csv_extra_columns_is_fatal(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "bad.csv"
            path.write_text("a,b\n1,2,3\n", encoding="utf-8")
            diagnostics = finalizer.Diagnostics(allow_incomplete=True)
            with self.assertRaisesRegex(
                finalizer.ReportFinalizationError, "열 수가 header보다 많습니다"
            ):
                finalizer._load_csv(path, diagnostics, "손상 테스트")

    def test_strict_numeric_cell_validation_rejects_nan_and_blank(self) -> None:
        diagnostics = finalizer.Diagnostics(allow_incomplete=False)
        finalizer._validate_csv_cells(
            [{"metric": "NaN"}, {"metric": ""}],
            numeric_columns=("metric",),
            diagnostics=diagnostics,
            label="수치 테스트",
        )
        with self.assertRaisesRegex(
            finalizer.ReportFinalizationError, "비수치/비유한값"
        ):
            diagnostics.fail_if_errors()

    def test_stale_embedded_path_is_a_clear_strict_error(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            diagnostics = finalizer.Diagnostics(allow_incomplete=False)
            finalizer._check_path_references(
                {"source_metadata": r"C:\obsolete\missing\run_metadata.json"},
                ("source_metadata",),
                Path(temporary),
                diagnostics,
                "경로 테스트",
            )
            with self.assertRaisesRegex(
                finalizer.ReportFinalizationError, "오래된 경로"
            ):
                diagnostics.fail_if_errors()

    def test_markdown_renderer_preserves_table_code_list_and_local_image(self) -> None:
        # A 1x1 transparent PNG.  Keeping the fixture inline avoids Pillow in the test.
        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            (root / "pixel.png").write_bytes(png)
            markdown = root / "sample.md"
            markdown.write_text(
                "# 제목\n\n본문 **강조**와 `코드`.\n\n"
                "- 목록\n\n| A | B |\n|---|---|\n| 1 | 2 |\n\n"
                "```python\nprint('ok')\n```\n\n![픽셀](pixel.png)\n",
                encoding="utf-8",
            )
            docx = root / "sample.docx"
            warnings: list[str] = []
            finalizer.markdown_to_docx(markdown, docx, warnings)
            self.assertTrue(docx.is_file())
            self.assertEqual(warnings, [])

            from docx import Document

            document = Document(docx)
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(document.tables[0].cell(1, 1).text, "2")
            all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("print('ok')", all_text)
            self.assertIn("목록", all_text)
            self.assertGreaterEqual(len(document.inline_shapes), 1)

    def test_verdict_count_and_paired_delta_use_only_recorded_rows(self) -> None:
        evidence = finalizer.Evidence(
            reproduction_verification=[
                {"verdict": "match"},
                {"verdict": "partial"},
                {"verdict": "match"},
            ],
            adaptive_repeats=[
                {
                    "mode": "fixed_paper_like",
                    "area": "300",
                    "policy": "Greedy",
                    "repeat": "0",
                    "finished": "10",
                    "avg_time": "5",
                    "semantic_quality": "0.9",
                },
                {
                    "mode": "adaptive_semantic",
                    "area": "300",
                    "policy": "Greedy",
                    "repeat": "0",
                    "finished": "13",
                    "avg_time": "4",
                    "semantic_quality": "0.88",
                },
            ],
        )
        self.assertEqual(finalizer._verdict_counts(evidence.reproduction_verification), {"match": 2, "partial": 1})
        paired = finalizer._paired_repeat_differences(evidence.adaptive_repeats)
        self.assertEqual(len(paired), 1)
        self.assertEqual(paired[0]["finished_delta"], 3.0)
        self.assertEqual(paired[0]["avg_time_delta"], -1.0)
        self.assertAlmostEqual(paired[0]["semantic_quality_delta"], -0.02)

    def test_current_feature_bitrate_and_nested_run_command_aliases(self) -> None:
        evidence = finalizer.Evidence(
            enhanced_airtalking={
                "schema_version": 2,
                "paper_like_active_channels": 80,
                "rho_c_feature_uncompressed_mean": 0.1041667,
                "rho_c_feature_zlib_mean": 0.0992756,
                "semantic_quality_miou_final": 0.305416,
                "pixel_accuracy_final": 0.828616,
                "rgb_reconstruction_psnr_db": 18.435,
                "rgb_reconstruction_ssim": 0.567506,
                "feature_encode_bitrate_mbps_median": 420.0577,
                "feature_decode_bitrate_mbps_median": 81.5977,
                "num_samples": 500,
                "multi_rate_profiles": [
                    {"active_channels": value}
                    for value in (20, 40, 60, 80, 120)
                ],
            }
        )
        rendered = finalizer._render_semantic_profile(evidence)
        self.assertIn("420.058 / 81.598 Mbps", rendered)

        command = "python airtalking_reproduction.py --workers 6"
        self.assertEqual(
            finalizer._recorded_command(
                {"run_provenance": {"command_windows": command}}
            ),
            command,
        )
        self.assertEqual(
            finalizer._recorded_command(
                {"run_provenance": {"argv": ["python", "runner.py"]}}
            ),
            "python runner.py",
        )
        self.assertNotIn("을(를)", finalizer._missing("검증 항목"))


if __name__ == "__main__":
    unittest.main()
