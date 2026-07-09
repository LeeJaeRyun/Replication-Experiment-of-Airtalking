from __future__ import annotations

import csv
import json
from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "outputs"
FIG_DIR = OUT_DIR / "figures"
QUALITY_CSV = OUT_DIR / "compression_quality.csv"
POLICY_JSON = OUT_DIR / "policy_summary.json"
REPORT_PATH = ROOT / "Adaptive_Semantic_Compression_Research_Proposal_KR.docx"

LATIN_FONT = "Calibri"
KOREAN_FONT = "Malgun Gothic"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "555555"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
POSITIVE_FILL = "EEF7F1"
CAUTION_FILL = "FFF8E8"
BORDER = "D7DBE2"


def require_inputs() -> None:
    required = [
        QUALITY_CSV,
        POLICY_JSON,
        FIG_DIR / "quality_vs_payload.png",
        FIG_DIR / "delivery_time_by_policy.png",
        FIG_DIR / "adaptive_mode_usage.png",
    ]
    missing = [str(path.relative_to(ROOT)) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required adaptive study outputs:\n" + "\n".join(missing))


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = LATIN_FONT
    style._element.get_or_add_rPr()
    rfonts = style._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        style._element.rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def paragraph_border_bottom(paragraph, color: str = BORDER, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_first_row_as_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(table, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, size: float = 8.7, bold: bool = False, color: str = "000000", align=None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.12
    if align is not None:
        para.alignment = align
    run = para.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    mark_first_row_as_header(table)
    set_cell_margins(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, header, size=font_size, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(values) > 2 else None
            set_cell_text(cells[idx], value, size=font_size, align=align)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    mark_first_row_as_header(table)
    set_cell_margins(table, top=150, bottom=150, start=190, end=190)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(title)
    set_run_font(r_title, size=10.5, color=NAVY, bold=True)
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    p_body.paragraph_format.line_spacing = 1.12
    r_body = p_body.add_run(body)
    set_run_font(r_body, size=9.8, color="222222")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_para(doc: Document, text: str, size: float = 10.6, color: str = "000000", bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=10.4)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=8.6, color=MUTED, italic=True)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.1) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image = p.add_run().add_picture(str(path), width=Inches(width))
    image._inline.docPr.set("descr", caption)
    image._inline.docPr.set("title", caption.split(".", 1)[0])
    add_caption(doc, caption)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    set_style_font(styles["Heading 1"], size=16, color=BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Heading 2"], size=13, color=BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    set_style_font(styles["List Bullet"], size=10.4)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Bullet"].paragraph_format.space_after = Pt(8)
    styles["List Bullet"].paragraph_format.line_spacing = 1.167
    header = section.header
    header.paragraphs[0].text = ""
    hp = header.paragraphs[0]
    hr = hp.add_run("Research proposal | Adaptive semantic compression")
    set_run_font(hr, size=8.5, color=MUTED)
    footer = section.footer
    footer.paragraphs[0].text = ""
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Generated: " + date.today().isoformat())
    set_run_font(fr, size=8.5, color=MUTED)


def load_quality_rows() -> list[dict[str, str]]:
    with QUALITY_CSV.open(newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def load_policy() -> dict:
    return json.loads(POLICY_JSON.read_text(encoding="utf-8"))


def fmt_pct(value: float, digits: int = 1) -> str:
    return f"{value * 100:.{digits}f}%"


def add_masthead(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("연구 제안서")
    set_run_font(r, size=23, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("채널 상태 기반 Adaptive Semantic Compression을 이용한 Multi-UAV D2D 통신 성능 개선")
    set_run_font(r, size=13.5, color=MUTED, bold=True)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["수신", "교수님"],
            ["작성 목적", "기존 AirTalking 논문의 한계를 바탕으로 후속 연구 주제 진행 가능 여부 검토 요청"],
            ["제안 주제", "Channel-aware Adaptive Semantic Compression for Multi-UAV D2D Communication"],
            ["핵심 아이디어", "통신 상태가 나쁘면 더 강하게 압축하고, 통신 상태가 좋으면 의미 품질을 더 보존하는 방식"],
            ["예비 실험", "Cityscapes label map 400개 기반 압축률-의미품질 trade-off 측정"],
        ],
        [1900, 7460],
        font_size=8.8,
    )
    add_callout(
        doc,
        "교수님께 여쭤보고 싶은 핵심 질문",
        "AirTalking 재현 결과를 기반으로, 고정된 semantic compression ratio 대신 채널 상태와 의미 품질 요구사항에 따라 compression level을 선택하는 연구 방향으로 진행해도 될지 검토받고자 합니다.",
        fill=POSITIVE_FILL,
    )


def write_background(doc: Document) -> None:
    doc.add_heading("1. 연구 배경", level=1)
    add_para(
        doc,
        "기존 AirTalking 논문은 여러 UAV가 지상 장치의 데이터를 중계하는 상황에서 semantic communication을 적용한다. "
        "원본 이미지를 그대로 보내는 대신, semantic encoder가 이미지의 의미 정보를 추출하고 decoder가 이를 다시 활용하면 전송해야 할 데이터 크기가 줄어든다.",
    )
    add_para(
        doc,
        "제가 수행한 재현 실험에서도 Cityscapes 공개 데이터로 semantic feature payload ratio를 측정했을 때 rho_c=0.104464가 나왔고, 논문값 rho_c=0.104와 거의 같았습니다. "
        "즉 '의미 정보만 보내면 payload가 크게 줄어든다'는 전제는 공개 데이터로도 확인 가능했습니다.",
    )
    add_callout(
        doc,
        "쉬운 설명",
        "사진 전체를 보내는 것은 원본 파일을 그대로 전송하는 것과 같습니다. semantic communication은 사진 전체 대신 '도로는 여기, 차는 여기, 사람은 여기' 같은 핵심 정보만 보내는 방식입니다. 인터넷이 느릴 때 고화질 영상 대신 낮은 화질로 보는 것과 비슷하게, 통신 상황에 맞춰 보내는 정보량을 조절할 수 있습니다.",
        fill=CALLOUT_FILL,
    )


def write_limitations(doc: Document) -> None:
    doc.add_heading("2. 기존 논문의 한계", level=1)
    add_table(
        doc,
        ["한계", "문제점", "후속 연구 방향"],
        [
            ["고정 semantic compression", "논문은 rho_c=0.104를 고정값처럼 사용한다. 채널이 좋거나 나쁜 상황을 구분하지 않는다.", "채널 상태에 따라 compression level을 다르게 선택한다."],
            ["semantic 품질 고려 부족", "payload가 줄어드는 효과는 보이지만, 의미 정보가 얼마나 정확히 보존되는지 scheduling에 직접 반영하지 않는다.", "mIoU, pixel accuracy 등 의미 품질 지표를 함께 최적화한다."],
            ["encoder/decoder 세부 비공개", "modified U-Net/Pix2PixHD라고 설명하지만 code, weight, feature map 정의가 공개되지 않았다.", "초기에는 Cityscapes label map proxy를 쓰고, 이후 공개 segmentation 모델로 확장한다."],
            ["시뮬레이션 파라미터 비공개", "request probability, workload, 전력값, 정책 hyperparameter가 명확하지 않아 exact reproduction이 어렵다.", "assumed parameter를 분리하고 민감도 분석을 수행한다."],
            ["소규모 환경 중심", "20 UAV/20 device 환경에서 검증되어 대규모 확장성은 충분히 확인되지 않았다.", "UAV/device 수 증가에 따른 성능 변화도 추가 실험한다."],
        ],
        [2150, 3550, 3660],
        font_size=7.8,
    )


def write_research_topic(doc: Document) -> None:
    doc.add_heading("3. 제안 연구 주제", level=1)
    add_para(
        doc,
        "제안 주제는 '채널 상태 기반 Adaptive Semantic Compression'입니다. 기존 방식은 semantic을 사용할 때 항상 같은 크기로 줄인다고 가정하지만, 제안 방식은 상황에 따라 압축 강도를 조절합니다.",
    )
    add_table(
        doc,
        ["상황", "선택 방식", "의도"],
        [
            ["채널 상태 나쁨", "emergency/low compression", "의미 품질을 조금 양보하더라도 데이터 크기를 줄여 전송 성공률과 지연시간을 개선"],
            ["채널 상태 중간", "medium/paper-like compression", "전송량과 의미 품질 사이의 균형 유지"],
            ["채널 상태 좋음", "high quality compression", "전송 여유가 있으므로 의미 품질을 더 보존"],
            ["배터리 부족 또는 요청 긴급", "더 강한 compression", "에너지와 시간을 아끼는 방향으로 선택"],
        ],
        [2200, 2800, 4360],
        font_size=8.2,
    )
    add_callout(
        doc,
        "연구 질문",
        "고정된 semantic compression ratio를 사용하는 것보다, 채널 상태와 semantic 품질 요구사항을 함께 고려해 compression level을 선택하면 UAV D2D 통신에서 완료 요청 수, 지연시간, 에너지 효율을 개선할 수 있는가?",
        fill=POSITIVE_FILL,
    )


def write_technical_plan(doc: Document) -> None:
    doc.add_heading("4. 기술적 구성", level=1)
    doc.add_heading("4.1 사용할 기술", level=2)
    add_table(
        doc,
        ["기술", "사용 위치", "비전공자용 설명"],
        [
            ["Computer Vision", "Cityscapes 이미지와 semantic label 처리", "사진 안에서 도로, 차, 사람 같은 의미를 다루는 기술"],
            ["Deep Learning", "semantic encoder 후보 모델 구성", "사진을 보고 의미 정보를 자동으로 뽑는 AI 모델"],
            ["Semantic Segmentation", "U-Net, DeepLabV3+, SegFormer 후보", "이미지의 각 픽셀이 어떤 물체인지 분류하는 작업"],
            ["Wireless Communication", "SINR, Shannon rate, UAV 링크 전송시간 계산", "신호가 좋으면 빠르고, 신호가 나쁘면 느린 현상을 수식으로 계산"],
            ["Scheduling/Optimization", "압축 단계와 UAV pair 선택", "어떤 드론이 어떤 방식으로 보낼지 결정하는 알고리즘"],
            ["Python Simulation", "전체 반복 실험과 결과 그래프 생성", "실제 드론 대신 컴퓨터 안에서 여러 조건을 반복 실험"],
        ],
        [1900, 3400, 4060],
        font_size=7.9,
    )

    doc.add_heading("4.2 연구 파이프라인", level=2)
    add_table(
        doc,
        ["단계", "내용", "산출물"],
        [
            ["1", "Cityscapes 이미지와 label map 준비", "원본 이미지, semantic label"],
            ["2", "semantic representation을 여러 compression level로 생성", "emergency/low/medium/paper-like/high payload"],
            ["3", "복원 품질 측정", "mIoU, pixel accuracy"],
            ["4", "UAV 통신 시뮬레이터에 compression profile 입력", "전송시간, 에너지, 처리 완료 수"],
            ["5", "adaptive controller 설계", "SINR, deadline, energy를 보고 compression level 선택"],
            ["6", "baseline과 비교", "raw, fixed semantic, adaptive semantic 성능 비교"],
        ],
        [900, 5100, 3360],
        font_size=8.1,
    )
    add_para(
        doc,
        "초기 단계에서는 딥러닝 모델 학습 없이 Cityscapes label map을 semantic feature proxy로 사용합니다. "
        "연구가 진행되면 공개 segmentation 모델(U-Net, DeepLabV3+, SegFormer 등)을 encoder로 사용해 실제 inference time과 semantic 품질을 측정하는 방향으로 확장할 수 있습니다.",
    )


def write_pilot_results(doc: Document, quality_rows: list[dict[str, str]], policy_data: dict) -> None:
    doc.add_heading("5. 예비 실험 결과", level=1)
    add_para(
        doc,
        "제안 주제가 실험적으로 말이 되는지 확인하기 위해 Cityscapes gtFine labelIds 400개를 사용해 간단한 파일럿 실험을 수행했습니다. "
        "label map을 여러 해상도로 줄인 뒤 다시 원래 크기로 복원하여, payload 크기와 의미 품질 사이의 trade-off를 측정했습니다.",
    )
    rows = []
    for row in quality_rows:
        rows.append(
            [
                row["mode"],
                f"{float(row['scale']):.3f}",
                fmt_pct(float(row["feature_ratio_mean"])),
                f"{float(row['mean_iou_mean']):.3f}",
                f"{float(row['pixel_accuracy_mean']):.3f}",
            ]
        )
    add_table(
        doc,
        ["Mode", "Scale", "Payload ratio", "mIoU", "Pixel acc."],
        rows,
        [1700, 1300, 1900, 1600, 2860],
        font_size=8.3,
    )
    add_figure(doc, FIG_DIR / "quality_vs_payload.png", "그림 1. Cityscapes label map 기반 compression-quality trade-off.", width=6.05)

    policy = policy_data["policy_summary"]
    add_table(
        doc,
        ["방식", "평균 전송시간", "평균 semantic 품질", "설명"],
        [
            ["Raw", f"{policy['raw_time_s']['mean']:.2f} s", "-", "원본 payload 전송"],
            ["Fixed paper-like", f"{policy['fixed_paper_time_s']['mean']:.2f} s", f"{policy['fixed_paper_miou']:.3f}", "AirTalking rho_c=0.104에 가까운 고정 압축"],
            ["Adaptive", f"{policy['adaptive_time_s']['mean']:.2f} s", f"{policy['adaptive_miou_mean']:.3f}", "SINR에 따라 compression level 선택"],
        ],
        [2100, 1900, 2000, 3360],
        font_size=8.2,
    )
    add_para(
        doc,
        f"단순 전송시간 추정에서는 adaptive 방식이 fixed paper-like 방식 대비 평균 전송시간을 약 {policy['adaptive_vs_fixed_paper_time_reduction_pct']:.1f}% 줄였습니다. "
        "다만 이 결과는 전체 UAV scheduling을 다시 돌린 최종 성능이 아니라, 연구 가능성을 보기 위한 파일럿 결과입니다.",
    )
    add_figure(doc, FIG_DIR / "delivery_time_by_policy.png", "그림 2. AirTalking 재현 실험에서 얻은 SINR sample 기반 평균 전송시간 추정.", width=6.05)
    add_figure(doc, FIG_DIR / "adaptive_mode_usage.png", "그림 3. 채널 상태에 따른 adaptive compression mode 선택 분포.", width=6.05)
    add_callout(
        doc,
        "예비 결과 해석",
        "고정 compression ratio 하나만 사용하는 것보다, 채널 상태에 따라 압축 강도를 바꾸는 연구 방향이 충분히 실험해볼 가치가 있음을 보여줍니다. 단, 최종 논문에서는 이 추정 결과를 그대로 주장하지 않고 UAV 시뮬레이터 전체에 adaptive policy를 넣어 다시 비교해야 합니다.",
        fill=CAUTION_FILL,
    )


def write_experiment_design(doc: Document) -> None:
    doc.add_heading("6. 본 실험 설계", level=1)
    add_table(
        doc,
        ["비교 대상", "설명"],
        [
            ["Baseline 1: Raw transmission", "원본 이미지 workload를 그대로 전송"],
            ["Baseline 2: Fixed semantic", "기존 AirTalking처럼 rho_c=0.104 수준의 고정 semantic compression 사용"],
            ["Baseline 3: Fixed low/high", "항상 낮은 품질 또는 항상 높은 품질 compression을 사용하는 단순 정책"],
            ["Proposed: Adaptive semantic", "SINR, deadline, battery, semantic quality target에 따라 compression level 선택"],
        ],
        [2850, 6510],
        font_size=8.3,
    )
    add_table(
        doc,
        ["평가 지표", "의미"],
        [
            ["Completed requests", "정해진 시간 안에 처리한 요청 수"],
            ["Average latency", "한 요청 처리에 걸린 평균 시간"],
            ["Flight/non-flight energy", "드론 이동/hover 에너지와 통신/encoding/decoding 에너지"],
            ["Semantic quality", "mIoU, pixel accuracy, task-level accuracy"],
            ["Scalability", "UAV/device 수 증가 시 계산량과 성능 변화"],
        ],
        [2850, 6510],
        font_size=8.3,
    )


def write_expected_contribution(doc: Document) -> None:
    doc.add_heading("7. 기대 기여와 진행 계획", level=1)
    add_bullet(doc, "기존 AirTalking의 고정 semantic compression 한계를 보완한다.")
    add_bullet(doc, "통신 성능뿐 아니라 semantic 품질을 함께 고려하는 scheduling 문제로 확장한다.")
    add_bullet(doc, "Cityscapes 기반 공개 데이터 실험으로 재현 가능성을 높인다.")
    add_bullet(doc, "학부 연구 수준에서는 label map proxy 기반 실험으로 시작하고, 시간이 허용되면 공개 segmentation encoder로 확장한다.")
    add_table(
        doc,
        ["기간", "작업"],
        [
            ["1주차", "관련 논문 정리, AirTalking 재현 코드 정리, adaptive compression level 정의"],
            ["2주차", "Cityscapes 기반 compression-quality profile 확장 측정"],
            ["3주차", "UAV simulator에 adaptive compression policy 추가"],
            ["4주차", "baseline 비교 실험 및 결과 그래프 생성"],
            ["5주차", "논문 초안 작성, 한계점과 추가 실험 정리"],
        ],
        [1600, 7760],
        font_size=8.4,
    )


def write_questions(doc: Document) -> None:
    doc.add_heading("8. 교수님께 확인받고 싶은 부분", level=1)
    add_bullet(doc, "이 주제를 학부생 학회 제출용 연구 주제로 잡아도 적절한지")
    add_bullet(doc, "딥러닝 encoder를 실제로 학습하는 수준까지 가야 하는지, 아니면 공개 모델 inference와 proxy 실험으로 충분한지")
    add_bullet(doc, "연구 범위를 adaptive compression policy에 집중할지, encoder 구현까지 포함할지")
    add_bullet(doc, "학회 제출을 위해 어떤 baseline과 metric을 반드시 넣어야 할지")
    add_callout(
        doc,
        "제안 결론",
        "본 연구는 AirTalking을 그대로 반복하는 것이 아니라, 기존 논문의 한계인 '고정 semantic compression'과 'semantic 품질 미반영'을 후속 연구 문제로 잡습니다. 구현 난이도는 조절 가능하고, 현재 보유한 Cityscapes 데이터와 재현 시뮬레이터를 그대로 활용할 수 있어 학부생 연구 주제로 현실성이 있습니다.",
        fill=POSITIVE_FILL,
    )


def write_references(doc: Document) -> None:
    doc.add_heading("참고 문헌 및 기반 자료", level=1)
    add_bullet(doc, "AirTalking: Aerial D2D for Multi-UAV Systems Based on Semantic Communication.")
    add_bullet(doc, "Cityscapes Dataset: urban scene understanding benchmark.")
    add_bullet(doc, "U-Net: Convolutional Networks for Biomedical Image Segmentation.")
    add_bullet(doc, "Pix2PixHD: High-Resolution Image Synthesis and Semantic Manipulation.")
    add_bullet(doc, "후속 encoder 후보: DeepLabV3+, SegFormer, lightweight U-Net 계열 공개 모델.")


def audit_docx(path: Path) -> dict[str, int | list[str]]:
    with ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    required = ["1. 연구 배경", "2. 기존 논문의 한계", "3. 제안 연구 주제", "4. 기술적 구성", "5. 예비 실험 결과"]
    return {
        "tables": xml.count("<w:tbl>"),
        "images": xml.count("<w:drawing>"),
        "header_rows": xml.count("<w:tblHeader"),
        "image_alt": xml.count("descr="),
        "required_sections_found": [title for title in required if title in xml],
    }


def build_docx() -> Path:
    require_inputs()
    quality_rows = load_quality_rows()
    policy_data = load_policy()
    doc = Document()
    configure_doc(doc)
    add_masthead(doc)
    write_background(doc)
    write_limitations(doc)
    write_research_topic(doc)
    write_technical_plan(doc)
    write_pilot_results(doc, quality_rows, policy_data)
    write_experiment_design(doc)
    write_expected_contribution(doc)
    write_questions(doc)
    write_references(doc)
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_docx()
    print(path)
    print(json.dumps(audit_docx(path), ensure_ascii=False, indent=2))
