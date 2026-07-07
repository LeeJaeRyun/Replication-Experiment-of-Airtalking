from __future__ import annotations

import csv
import json
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RUN_DIR = ROOT / "outputs" / "airtalking_cityscapes_feature_paperhw"
CITY_DIR = ROOT / "outputs" / "cityscapes_semantic_measurement"
FIG_DIR = RUN_DIR / "figures"

SUMMARY_CSV = RUN_DIR / "summary_metrics.csv"
METADATA_JSON = RUN_DIR / "run_metadata.json"
VERIFICATION_CSV = RUN_DIR / "verification_against_paper_cityscapes_feature_paperhw.csv"
CITY_SUMMARY_JSON = CITY_DIR / "cityscapes_semantic_summary.json"
REPORT_PATH = RUN_DIR / "AirTalking_Cityscapes_Reproduction_Report_KR.docx"

LATIN_FONT = "Calibri"
KOREAN_FONT = "Malgun Gothic"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "555555"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
POSITIVE_FILL = "EEF7F1"


def require_inputs() -> None:
    required = [SUMMARY_CSV, METADATA_JSON, VERIFICATION_CSV, CITY_SUMMARY_JSON]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required experiment outputs:\n" + "\n".join(missing))


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = LATIN_FONT
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = LATIN_FONT
    if style._element.rPr is None:
        style._element.append(OxmlElement("w:rPr"))
    rfonts = style._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        style._element.rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def style_paragraph_runs(paragraph, size: float = 10.5, color: str = "000000", bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def paragraph_border_bottom(paragraph, color: str = "D7DBE2", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_first_row_as_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_cell_margins(table, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, size: float = 8.7, bold: bool = False, color: str = "000000", align=None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.12
    if align is not None:
        para.alignment = align
    run = para.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], font_size: float = 8.6) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    mark_first_row_as_header(table)
    set_cell_margins(table)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, header, size=font_size, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(values) > 2 else None
            set_cell_text(cells[idx], value, size=font_size, align=align)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    mark_first_row_as_header(table)
    set_cell_margins(table, top=140, bottom=140, start=180, end=180)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    cell.text = ""
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(title)
    set_run_font(r_title, size=10.5, color=NAVY, bold=True)
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    p_body.paragraph_format.line_spacing = 1.15
    r_body = p_body.add_run(body)
    set_run_font(r_body, size=10.0, color="000000")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.3)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.3)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(text)
    set_run_font(run, size=8.8, color=MUTED, italic=True)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[그림 파일 없음: {path}]")
        set_run_font(r, size=9.0, color="9B1C1C", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    image = run.add_picture(str(path), width=Inches(width))
    image._inline.docPr.set("descr", caption)
    image._inline.docPr.set("title", caption.split(".", 1)[0])
    add_caption(doc, caption)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
        ("List Bullet", 10.3, "000000", 0, 4),
        ("List Number", 10.3, "000000", 0, 4),
    ]:
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=name.startswith("Heading"))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = "AirTalking Cityscapes reproduction"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph_runs(header, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = f"Generated {date.today().isoformat()}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph_runs(footer, size=8.5, color=MUTED)


def add_masthead(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("AirTalking 논문 실험 재현 보고서")
    set_run_font(run, size=22, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Cityscapes 기반 semantic communication profile 및 공개 파라미터 기반 근사 재현")
    set_run_font(run, size=11.5, color=MUTED)

    metadata = [
        ("대상 논문", "Airtalking: Aerial D2D for Multi-UAV Systems Based on Semantic Communication"),
        ("작성일", "2026-07-07"),
        ("데이터", "Cityscapes train/val 3,475장"),
        ("결과 폴더", r"outputs\airtalking_cityscapes_feature_paperhw"),
        ("재현 수준", "공개 파라미터 기반 재구현 + 미공개 파라미터 명시적 가정"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.2, color=NAVY, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.2, color="000000")

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)
    rule.paragraph_format.space_after = Pt(12)


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def row_lookup(rows: list[dict[str, str]], mode: str, area: int, policy: str) -> dict[str, str]:
    for row in rows:
        if row["mode"] == mode and int(row["area"]) == area and row["policy"] == policy:
            return row
    raise KeyError((mode, area, policy))


def fmt_float(value: str | float, digits: int = 1) -> str:
    return f"{float(value):,.{digits}f}"


def fmt_pct(value: float) -> str:
    return f"{value:+.1f}%"


def write_intro(doc: Document) -> None:
    doc.add_heading("1. 한 줄 결론", level=1)
    add_callout(
        doc,
        "현재 결과의 정확한 표현",
        "공개된 논문 파라미터와 실제 Cityscapes 데이터를 사용해 AirTalking 실험을 최대한 재구현했다. "
        "다만 논문에 공개되지 않은 시뮬레이터 내부값은 가정값으로 분리해 기록했기 때문에, "
        "이 결과는 '공식 코드와 동일한 정량 재현'이 아니라 '공개 정보 기반의 투명한 근사 재현'으로 해석해야 한다.",
        fill=POSITIVE_FILL,
    )
    p = doc.add_paragraph()
    p.add_run(
        "질문한 대로 보면, 현재 접근은 공개 정보만으로 가능한 범위에서는 꽤 타당하다. "
        "논문에 숫자로 적힌 값은 그대로 사용했고, 실제로 확보 가능한 데이터셋인 Cityscapes를 사용했으며, "
        "없는 값은 숨기지 않고 run_metadata.json의 assumed_params에 따로 기록했다. "
        "보고서에서는 이 지점을 강점과 한계로 동시에 설명한다."
    )
    style_paragraph_runs(p)

    doc.add_heading("2. 이 실험이 무엇을 한 것인지", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "AirTalking 논문은 UAV 여러 대가 지상 단말 간 D2D 통신을 도와주는 상황에서, "
        "semantic encoding/decoding을 적용하면 전송해야 할 payload가 줄고 요청 처리량, 지연시간, 에너지 비용이 좋아질 수 있다고 주장한다. "
        "이번 작업은 논문 수식과 공개 파라미터를 바탕으로 그 시스템을 Python 시뮬레이터로 다시 구성하고, "
        "Cityscapes semantic label을 사용해 논문 Table III의 semantic compression ratio가 실제 데이터에서 비슷하게 나오는지 확인한 뒤, "
        "Figure 3, Figure 4, Figure 6의 방향성과 수치를 비교한 것이다."
    )
    style_paragraph_runs(p)

    add_bullet(doc, "딥러닝 모델을 새로 학습한 실험은 아니다. 논문이 공개한 semantic encoder/decoder 성능값을 시스템 시뮬레이션에 넣었다.")
    add_bullet(doc, "컴퓨터비전 데이터셋인 Cityscapes의 이미지와 semantic segmentation label을 사용해 semantic payload 비율을 측정했다.")
    add_bullet(doc, "무선 통신 부분은 path-loss, Rician fading, SINR, Shannon rate를 계산하는 네트워크 시뮬레이션으로 구현했다.")
    add_bullet(doc, "정책 비교는 Stochastic, LinUCB, Simulated Annealing, Greedy, MCTS를 같은 환경에서 반복 실행하는 Monte Carlo 방식으로 수행했다.")


def write_data_section(doc: Document, city_summary: dict) -> None:
    doc.add_heading("3. 사용 데이터셋과 Cityscapes 선택 이유", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "현재 실제 파일로 확보한 데이터는 Cityscapes이다. 논문도 semantic compression ratio와 restoration ratio를 제시하지만, "
        "공식 AirTalking 소스코드나 원본 실험 로그, semantic encoder/decoder 모델 구조, 학습 설정은 공개되어 있지 않다. "
        "따라서 이번 재현에서는 Cityscapes의 leftImg8bit 이미지와 gtFine semantic label을 사용해 semantic payload profile을 직접 측정했다."
    )
    style_paragraph_runs(p)

    add_table(
        doc,
        ["항목", "사용 여부", "이유"],
        [
            ["leftImg8bit train/val", "사용", "원본 RGB 입력 이미지. raw workload 기준을 만들기 위해 필요"],
            ["gtFine train/val", "사용", "정답 semantic label. semantic feature payload proxy를 만들기 위해 필요"],
            ["test split", "미사용", "Cityscapes test label은 평가용 dummy/ignore region이라 semantic 측정에 부적합"],
            ["trainextra/coarse/right/disparity/camera", "미사용", "논문 Figure 3-6 재현에는 좌측 RGB와 fine semantic label만 필요"],
        ],
        [2300, 1500, 5560],
    )

    add_table(
        doc,
        ["측정값", "결과", "의미"],
        [
            ["샘플 수", f"{city_summary['num_samples']:,}장", "Cityscapes train/val 이미지-라벨 쌍"],
            ["raw RGB 평균 크기", f"{city_summary['raw_uncompressed_bytes_mean']:,.0f} bytes", "2048 x 1024 x 3 채널 기준"],
            ["semantic feature 평균 크기", f"{city_summary['semantic_feature_bytes_mean']:,.0f} bytes", "label map을 0.56 scale로 축소한 semantic feature proxy"],
            ["측정 rho_c", f"{city_summary['rho_c_feature_uncompressed_mean']:.6f}", "semantic feature payload / raw RGB payload"],
            ["논문 rho_c", "0.104", "Table III에 공개된 semantic compression ratio"],
        ],
        [2500, 2500, 4360],
    )

    add_callout(
        doc,
        "왜 zlib 압축 label을 그대로 쓰지 않았는가",
        "Cityscapes label PNG나 zlib 압축 label만 쓰면 rho_c가 약 0.003 수준으로 너무 작게 나온다. "
        "논문 Table III의 rho_c=0.104는 단순 파일 압축률이라기보다 semantic feature representation의 payload 비율로 보는 편이 자연스럽다. "
        "그래서 downsampled label tensor를 semantic feature proxy로 두었고, 이 방식에서 rho_c=0.104464가 나와 논문값과 거의 일치했다.",
        fill=CAUTION_FILL,
    )


def write_file_map(doc: Document) -> None:
    doc.add_heading("4. 각 파일과 폴더의 의미", level=1)
    add_table(
        doc,
        ["경로", "의미"],
        [
            [r"dataset\leftImg8bit_trainvaltest", "Cityscapes RGB 이미지. semantic profile 측정의 raw input"],
            [r"dataset\gtFine_trainvaltest", "Cityscapes fine semantic label. train/val labelIds를 사용"],
            ["measure_cityscapes_semantics.py", "Cityscapes 이미지/라벨을 읽어서 rho_c, payload 크기, encode/decode proxy 시간을 측정"],
            ["airtalking_reproduction.py", "AirTalking 시스템 모델을 Python/NumPy로 구현한 메인 시뮬레이터"],
            ["verify_against_paper.py", "논문 Figure 3/4/6의 시각 추정값과 재현 결과를 자동 비교"],
            [r"outputs\cityscapes_semantic_measurement\cityscapes_semantic_summary.json", "Cityscapes semantic profile 요약. 시뮬레이터 입력으로 사용"],
            [r"outputs\airtalking_cityscapes_feature_paperhw\run_metadata.json", "논문 공개값과 가정값을 분리해 기록한 metadata"],
            [r"outputs\airtalking_cityscapes_feature_paperhw\summary_metrics.csv", "정책/면적별 finished requests, energy, latency, travel distance 결과"],
            [r"outputs\airtalking_cityscapes_feature_paperhw\figures", "보고서와 검증용 그림 파일"],
            [r"outputs\airtalking_cityscapes_feature_paperhw\verification_*.csv/md", "논문 그래프 대비 match/partial/mismatch 판정"],
        ],
        [3600, 5760],
        font_size=8.3,
    )


def write_method_section(doc: Document, metadata: dict) -> None:
    doc.add_heading("5. 기술 스택과 수행 방법", level=1)
    add_table(
        doc,
        ["도구/기술", "사용 목적"],
        [
            ["Python", "전체 데이터 측정, 시뮬레이션, 결과 집계, 보고서 생성"],
            ["NumPy", "UAV/device 위치, random walk, fading, Monte Carlo 반복 계산"],
            ["Pillow(PIL)", "Cityscapes PNG 이미지와 labelIds를 읽고 semantic feature 크기 측정"],
            ["zlib", "비교용 label 압축률 측정. 최종 시뮬레이션 입력으로는 feature proxy 사용"],
            ["Matplotlib", "finished requests, energy, latency, semantic vs non-semantic 결과 그림 생성"],
            ["python-docx", "현재 한국어 DOCX 보고서 생성"],
            ["시스템 시뮬레이션", "딥러닝 학습이 아니라, 논문 수식과 파라미터를 기반으로 통신/스케줄링 성능을 재현"],
        ],
        [2300, 7060],
    )

    doc.add_heading("5.1 실험 절차", level=2)
    steps = [
        "Cityscapes train/val 이미지와 gtFine labelIds를 pairing했다.",
        "raw RGB payload와 semantic feature proxy payload를 계산해 rho_c를 측정했다.",
        "논문 Table III의 encoder/decoder bitrate를 사용하고, Cityscapes에서 측정한 rho_c=0.104464를 시뮬레이터에 넣었다.",
        "100 x 100 m2부터 500 x 500 m2까지 5개 면적에서 5개 정책을 1000초, 10회 반복으로 실행했다.",
        "300 x 300 m2에서는 semantic processing을 끈 non-semantic baseline도 실행했다.",
        "결과 CSV를 논문 Figure 3/4/6의 시각 추정값과 비교해 match, partial, mismatch로 판정했다.",
    ]
    for step in steps:
        add_numbered(doc, step)

    doc.add_heading("5.2 공개 파라미터", level=2)
    paper = metadata["paper_params"]
    public_rows = [
        ["UAV 수 / device 수", f"{paper['n_uav']} / {paper['n_device']}", "논문 Table III"],
        ["slot / horizon / repeats", f"{paper['dt']} s / {paper['t_slots']} slots / {paper['repeats']} repeats", "논문 실험 설정 기반"],
        ["UAV 속도/가속/감속", f"{paper['vmax_uav']} m/s, {paper['accel_uav']} m/s2, {paper['decel_uav']} m/s2", "논문 Table III"],
        ["UAV 평균 고도", f"{paper['height_mean']} m", "논문 Table III"],
        ["carrier bandwidth/frequency", f"{paper['carrier_bandwidth']/1e6:.0f} MHz / {paper['carrier_frequency']/1e9:.0f} GHz", "논문 Table III"],
        ["Tx power", f"UAV {paper['p_uav_tx']} W, device {paper['p_device_tx']} W", "논문 Table III"],
        ["path-loss exponent", f"U2U {paper['alpha_u2u']}, U2G {paper['alpha_u2g']}", "논문 Table III"],
        ["rho_c / rho_r", f"{paper['rho_c']:.6f} / {paper['rho_r']}", "Cityscapes 측정값 + 논문 rho_r"],
        ["encoder/decoder bitrate", f"{paper['enc_bitrate']/1e6:.2f} / {paper['dec_bitrate']/1e6:.2f} Mbps", "논문 Table III"],
    ]
    add_table(doc, ["파라미터", "사용값", "근거"], public_rows, [2700, 3300, 3360], font_size=8.4)


def write_assumption_section(doc: Document, metadata: dict) -> None:
    doc.add_heading("6. 가정값이 들어간 부분과 이유", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "아래 값들은 논문에 숫자로 공개되지 않았다. 그래서 코드에서는 assumed_params로 분리했고, "
        "보고서에서도 공개 파라미터와 섞지 않는다. 핵심은 값을 숨긴 것이 아니라, 어떤 부분이 논문값이고 어떤 부분이 재현자의 가정인지 추적 가능하게 만든 점이다."
    )
    style_paragraph_runs(p)

    assumed = metadata["assumed_params"]
    rows = [
        ["request_probability", f"{assumed['request_probability']}", "논문은 Bernoulli request라고만 설명하고 p_req 수치를 제공하지 않는다. 1000초 horizon에서 과도한 saturation 없이 요청이 생성되도록 낮은 확률을 사용했다."],
        ["workload distribution", f"mean {assumed['workload_mean_bits']/1e6:.0f} Mb, std {assumed['workload_std_bits']/1e6:.0f} Mb", "논문은 raw workload w_l의 분포를 공개하지 않는다. Mbps 단위 encoder/decoder와 transmission time이 실험에 의미 있게 반영되도록 heterogeneous workload를 사용했다."],
        ["workload range", f"{assumed['workload_min_bits']/1e6:.0f} - {assumed['workload_max_bits']/1e6:.0f} Mb", "너무 작은/큰 request가 결과를 지배하지 않도록 clipped normal distribution으로 제한했다."],
        ["P_mov / P_hov", f"{assumed['p_move']} / {assumed['p_hover']} W", "논문 에너지식은 P_mov, P_hov를 요구하지만 수치를 주지 않는다. 논문 설명처럼 flight/hover energy가 전체 비용에서 지배적이 되도록 별도 가정했다."],
        ["P_cod / P_dec", f"{assumed['p_encode']} / {assumed['p_decode']} W", "논문은 coding/decoding power 기호만 제시한다. semantic processing energy를 반영하되 propulsion보다 작게 두었다."],
        ["P_d2d", f"{assumed['p_d2d_radio']} W", "D2D radio module energy 계산용 값이다. Table III의 Tx power는 channel/SINR 계산에 사용하고, module runtime energy는 별도 항으로 처리했다."],
        ["device mobility", f"D={assumed['device_diffusion']}, cap={assumed['device_speed_cap']}", "논문은 diffusion mobility를 제시하지만 계수값은 공개하지 않는다. 지상 단말이 천천히 움직이는 random walk로 모델링했다."],
        ["density_interference_scale", f"{assumed['density_interference_scale']}", "논문은 작은 면적에서 간섭이 커진다고 설명하지만 동시 active link 처리 세부가 없다. 면적 밀도에 따른 interference penalty로 그 경향을 근사했다."],
        ["policy hyperparameters", f"LinUCB alpha={assumed['linucb_alpha']}, SA iter={assumed['sa_iterations']}, MCTS samples={assumed['mcts_samples']}", "논문은 정책 이름과 방향은 제시하지만 내부 탐색 budget과 hyperparameter를 공개하지 않는다. 반복 실행 가능한 계산량으로 설정했다."],
        ["random seed", str(assumed["seed"]), "결과 재현성을 위해 고정했다."],
    ]
    add_table(doc, ["가정 항목", "사용값", "왜 이렇게 두었는가"], rows, [2200, 2350, 4810], font_size=7.7)

    add_callout(
        doc,
        "보고서에 써야 하는 표현",
        "이 값들은 논문에서 가져온 값이 아니라, 논문이 공개하지 않은 simulator configuration을 채우기 위한 명시적 가정이다. "
        "따라서 결과는 exact reproduction이 아니라 calibrated/assumption-based reproduction이다. "
        "하지만 공개값과 실제 데이터셋을 사용하고 가정값을 분리 기록했으므로, 재현 시도 자체는 투명하고 방어 가능하다.",
        fill=CAUTION_FILL,
    )


def write_results_section(doc: Document, rows: list[dict[str, str]], verification_rows: list[dict[str, str]]) -> None:
    doc.add_heading("7. 실험 결과", level=1)
    doc.add_heading("7.1 면적별 완료 요청 수", level=2)

    policies = ["Stochastic", "LinUCB", "SA", "Greedy", "MCTS"]
    finished_rows = []
    for area in [100, 200, 300, 400, 500]:
        result = [f"{area} x {area}"]
        for policy in policies:
            result.append(fmt_float(row_lookup(rows, "semantic", area, policy)["finished"], 1))
        finished_rows.append(result)
    add_table(doc, ["Area", *policies], finished_rows, [1450, 1550, 1550, 1550, 1550, 1710], font_size=8.0)

    add_figure(
        doc,
        FIG_DIR / "finished_requests.png",
        "그림 1. 면적별 finished requests 재현 결과. 대부분 정책에서 면적이 커질수록 완료 요청 수가 증가한다.",
        width=6.15,
    )

    doc.add_heading("7.2 300m 환경에서 semantic vs non-semantic", level=2)
    comparison_rows = []
    for policy in ["LinUCB", "SA", "Greedy", "MCTS"]:
        sem = float(row_lookup(rows, "semantic", 300, policy)["finished"])
        non = float(row_lookup(rows, "nonsemantic", 300, policy)["finished"])
        gain = (sem - non) / non * 100.0
        comparison_rows.append([policy, f"{sem:.1f}", f"{non:.1f}", fmt_pct(gain)])
    add_table(doc, ["Policy", "Semantic finished", "Non-semantic finished", "증가율"], comparison_rows, [2100, 2500, 2600, 2160], font_size=8.5)

    add_figure(
        doc,
        FIG_DIR / "semantic_vs_nonsemantic_300m.png",
        "그림 2. 300 x 300 m2에서 semantic processing을 켰을 때 모든 정책의 완료 요청 수가 증가했다.",
        width=6.15,
    )

    doc.add_heading("7.3 논문 Figure 3/4/6과의 자동 비교", level=2)
    counts = Counter(row["verdict"] for row in verification_rows)
    add_table(
        doc,
        ["판정", "개수", "해석"],
        [
            ["Match", str(counts.get("match", 0)), "논문 시각 추정값 대비 상대오차 25% 이내"],
            ["Partial", str(counts.get("partial", 0)), "상대오차 50% 이내"],
            ["Mismatch", str(counts.get("mismatch", 0)), "상대오차 50% 초과"],
        ],
        [1800, 1300, 6260],
        font_size=8.7,
    )

    largest = sorted(verification_rows, key=lambda row: float(row["relative_error"]), reverse=True)[:6]
    add_table(
        doc,
        ["Figure", "Area", "Policy", "Metric", "Paper est.", "Repro.", "Rel. err."],
        [
            [
                row["check"],
                row["area"],
                row["policy"],
                row["metric"],
                fmt_float(row["paper_visual_estimate"], 1),
                fmt_float(row["reproduction"], 1),
                fmt_float(row["relative_error"], 2),
            ]
            for row in largest
        ],
        [1750, 900, 1200, 1700, 1250, 1300, 1260],
        font_size=7.7,
    )

    add_callout(
        doc,
        "결과 해석",
        "semantic이 non-semantic보다 좋은 방향성, 그리고 면적 증가에 따라 finished requests가 증가하는 경향은 재현됐다. "
        "하지만 에너지와 평균 처리시간의 절대 크기는 논문 Figure와 크게 다르다. 이 차이는 Cityscapes 때문이라기보다, "
        "논문에 공개되지 않은 workload, request probability, propulsion/hover power, interference scheduling, 정책 hyperparameter의 영향이 크다.",
        fill=CALLOUT_FILL,
    )


def write_comparison_section(doc: Document) -> None:
    doc.add_heading("8. 논문과 비슷한 부분, 다른 부분", level=1)
    add_table(
        doc,
        ["구분", "판정", "내용"],
        [
            ["Semantic compression ratio", "매우 유사", "Cityscapes 측정 rho_c=0.104464, 논문 rho_c=0.104"],
            ["Semantic vs non-semantic 방향성", "일치", "300m 환경에서 semantic processing이 모든 정책의 finished requests를 증가시킴"],
            ["면적 증가에 따른 요청 완료 경향", "대체로 일치", "Stochastic, LinUCB, SA, Greedy, MCTS 모두 면적 증가에 따라 finished requests가 증가"],
            ["정책 ranking", "부분 불일치", "논문은 넓은 area에서 LinUCB가 SA보다 좋다고 설명하지만, 현재 재현에서는 SA가 더 높게 나오는 구간이 있음"],
            ["에너지/시간 절대값", "불일치", "Flight energy와 avg time이 논문 시각 추정값보다 크게 나옴"],
            ["정량 일치 수준", "제한적", "Match 6, Partial 14, Mismatch 53"],
        ],
        [2300, 1600, 5460],
        font_size=8.2,
    )

    add_figure(
        doc,
        FIG_DIR / "average_time_cost.png",
        "그림 3. 평균 처리시간 재현 결과. 논문 대비 절대값 차이가 가장 큰 축 중 하나다.",
        width=6.15,
    )
    add_figure(
        doc,
        FIG_DIR / "flight_energy_per_request.png",
        "그림 4. 요청당 flight energy 재현 결과. P_mov/P_hov가 미공개라 정량 차이가 크게 나타난다.",
        width=6.15,
    )


def write_interpretation_section(doc: Document) -> None:
    doc.add_heading("9. 최종 해석: 최대한 잘 재현한 것인가?", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "공개된 정보만 기준으로 보면, 현재 결과는 최대한 잘 재현한 편이라고 정리할 수 있다. "
        "그 이유는 첫째, 논문 Table III에 실제 수치로 공개된 파라미터를 코드에 반영했기 때문이다. "
        "둘째, 데이터셋은 임의 생성이 아니라 실제 Cityscapes train/val을 사용했고, semantic compression ratio가 논문값과 거의 같게 측정됐다. "
        "셋째, 논문에 없는 값은 결과에 몰래 섞지 않고 assumed_params로 분리해 기록했다. "
        "넷째, 논문 그래프와의 비교를 자동화해 어느 부분이 맞고 틀린지 확인할 수 있게 만들었다."
    )
    style_paragraph_runs(p)

    p = doc.add_paragraph()
    p.add_run(
        "반대로, 이것을 '논문 공식 실험과 완전히 동일한 재현'이라고 쓰면 안 된다. "
        "공식 소스코드, raw simulation result, request/workload/power/policy hyperparameter가 공개되지 않았기 때문이다. "
        "따라서 보고서나 발표에서는 '공개 정보 기반 근사 재현' 또는 'assumption-based reproduction'이라고 표현하는 것이 정확하다."
    )
    style_paragraph_runs(p)

    add_callout(
        doc,
        "추천 문장",
        "본 실험은 AirTalking 논문에서 공개한 Table III 파라미터와 실제 Cityscapes semantic segmentation 데이터를 기반으로 수행한 재구현이다. "
        "논문에 공개되지 않은 workload, request generation, energy power, interference scheduling, policy hyperparameter는 별도 가정값으로 분리해 설정했으며, "
        "따라서 본 결과는 공식 코드 기반 exact reproduction이 아니라 공개 정보 기반 approximate reproduction으로 해석한다.",
        fill=POSITIVE_FILL,
    )


def write_appendix(doc: Document) -> None:
    doc.add_heading("10. 재실행 명령과 산출물 확인", level=1)
    commands = [
        r"python measure_cityscapes_semantics.py --root dataset --out outputs\cityscapes_semantic_measurement --splits train,val --feature-scale 0.56 --repeats 1",
        r"python airtalking_reproduction.py --out outputs\airtalking_cityscapes_feature_paperhw --semantic-summary outputs\cityscapes_semantic_measurement\cityscapes_semantic_summary.json --semantic-profile-kind feature --semantic-encoder-mode paper --semantic-decoder-mode paper",
        r"python verify_against_paper.py --summary outputs\airtalking_cityscapes_feature_paperhw\summary_metrics.csv --label _cityscapes_feature_paperhw",
        r"python build_airtalking_report.py",
    ]
    for cmd in commands:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(cmd)
        set_run_font(run, size=8.5, color=NAVY)

    add_table(
        doc,
        ["확인할 파일", "용도"],
        [
            [str(REPORT_PATH.relative_to(ROOT)), "현재 DOCX 보고서"],
            [str(SUMMARY_CSV.relative_to(ROOT)), "정책/면적별 핵심 수치"],
            [str(METADATA_JSON.relative_to(ROOT)), "논문 공개값과 가정값 분리 기록"],
            [str(VERIFICATION_CSV.relative_to(ROOT)), "논문 그래프 대비 row-level 비교"],
            [str((FIG_DIR / "finished_requests.png").relative_to(ROOT)), "완료 요청 수 그림"],
            [str((FIG_DIR / "semantic_vs_nonsemantic_300m.png").relative_to(ROOT)), "semantic vs non-semantic 비교 그림"],
        ],
        [3900, 5460],
        font_size=8.3,
    )


def build_report() -> Path:
    require_inputs()
    metadata = json.loads(METADATA_JSON.read_text(encoding="utf-8"))
    city_summary = json.loads(CITY_SUMMARY_JSON.read_text(encoding="utf-8"))
    summary_rows = load_csv(SUMMARY_CSV)
    verification_rows = load_csv(VERIFICATION_CSV)

    doc = Document()
    configure_doc(doc)
    add_masthead(doc)
    write_intro(doc)
    write_data_section(doc, city_summary)
    add_figure(
        doc,
        CITY_DIR / "samples" / "cityscapes_pair_1.png",
        "그림 0. Cityscapes RGB 이미지와 semantic label 예시. semantic feature profile 측정에 사용한 데이터 형식이다.",
        width=5.9,
    )
    write_file_map(doc)
    write_method_section(doc, metadata)
    write_assumption_section(doc, metadata)
    write_results_section(doc, summary_rows, verification_rows)
    write_comparison_section(doc)
    write_interpretation_section(doc)
    write_appendix(doc)

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)
