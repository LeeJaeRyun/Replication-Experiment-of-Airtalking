from __future__ import annotations

"""Readable, provenance-aware DOCX rendering for the research reports.

The Markdown remains the canonical research record.  This module is only the
presentation layer: it gives every generated DOCX the same cover, navigation,
typography, table treatment, figure captions, and source notes.
"""

import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Sequence
from xml.etree import ElementTree


ROOT = Path(__file__).resolve().parents[1]

NAVY = "17324D"
NAVY_2 = "244A68"
TEAL = "16858C"
TEAL_DARK = "0F6970"
CORAL = "E66A4E"
INK = "24313D"
MUTED = "5E6C78"
PALE_BLUE = "EAF3F6"
PALE_TEAL = "E8F4F3"
PALE_CORAL = "FCEDE8"
PALE_YELLOW = "FFF7DE"
PAPER = "F6F8FA"
RULE = "D8E1E7"
WHITE = "FFFFFF"

BODY_FONT = "Malgun Gothic"
MONO_FONT = "Consolas"

CITYSCAPES_CITATION_URL = "https://www.cityscapes-dataset.com/citation/"

FIGURE_PROVENANCE: dict[str, str] = {
    "studies/neural_encoder_decoder/results/enhanced_scalable_full_256x128_verified/qualitative_panel_paper_like.png": "cityscapes_panel",
    "studies/airtalking_reproduction/results/airtalking_enhanced_scalable_verified/figures/finished_requests.png": "airtalking_plot",
    "studies/airtalking_reproduction/results/airtalking_enhanced_scalable_verified/figures/semantic_vs_nonsemantic_300m.png": "airtalking_plot",
    "studies/adaptive_semantic_compression/results/adaptive_enhanced_scalable_verified/figures/finished_by_area_greedy_mcts.png": "adaptive_plot",
    "studies/adaptive_semantic_compression/results/adaptive_enhanced_scalable_verified/figures/adaptive_mode_usage_300m.png": "adaptive_plot",
    "studies/adaptive_semantic_compression/results/adaptive_enhanced_scalable_verified/figures/latency_quality_tradeoff_300m.png": "adaptive_plot",
}


REPORT_PRESENTATION: dict[str, dict[str, str]] = {
    "01": {
        "kicker": "RESEARCH REPORT 01  ·  NEURAL CODEC",
        "subtitle": "의미통신용 신경망 코덱의 학습 과정, 성능, 한계를 한 흐름으로 읽는 기록",
        "focus": "학습 과정 → rate–quality 결과 → 재현성·한계 순서로 읽으면 핵심이 빠르게 잡힙니다.",
        "short": "신경망 인코더·디코더",
        "chip": "VERIFIED RESULTS   ●   LOCAL GPU EXPERIMENT   ●   RTX 4060 Ti",
    },
    "02": {
        "kicker": "REPLICATION REPORT 02  ·  AIR TALKING",
        "subtitle": "논문 수치, 그림 판독값, 독립 구현 결과를 구분해 정리한 재현 기록",
        "focus": "‘논문값 / 그림 판독 근삿값 / 독립 구현값’의 출처 구분을 먼저 확인하세요.",
        "short": "AirTalking 실험 재현",
        "chip": "REPLICATION STUDY   ●   LOCAL SIMULATION   ●   RTX-DERIVED CODEC",
    },
    "03": {
        "kicker": "FOLLOW-UP STUDY 03  ·  ADAPTIVE COMPRESSION",
        "subtitle": "고정 압축과 채널 적응형 압축의 지연·품질 균형을 검증한 후속 연구",
        "focus": "평균값만 보지 말고 paired 변화, 95% 신뢰구간, guardrail을 함께 읽으세요.",
        "short": "적응형 의미 압축",
        "chip": "FOLLOW-UP STUDY   ●   LOCAL SIMULATION   ●   RTX-DERIVED PROFILE",
    },
    "04": {
        "kicker": "COMPARISON BRIEF 04  ·  BEFORE & AFTER",
        "subtitle": "이전·현재 컴퓨터와 서로 다른 실험 설계의 차이를 분리해 읽는 비교 기록",
        "focus": "하드웨어 차이와 데이터·모델·평가 설계 차이를 같은 원인으로 해석하지 않도록 구성했습니다.",
        "short": "이전·현재 결과 비교",
        "chip": "COMPARISON BRIEF   ●   PREVIOUS CPU   ●   CURRENT RTX 4060 Ti",
    },
}


def _rgb(hex_color: str) -> Any:
    from docx.shared import RGBColor

    return RGBColor.from_string(hex_color)


def _set_font(
    run: Any,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    monospace: bool = False,
    family: str | None = None,
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    selected_family = family or (MONO_FONT if monospace else BODY_FONT)
    run.font.name = selected_family
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), selected_family)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), selected_family)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), selected_family)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = _rgb(color)


def _set_shading(element: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    old = properties.find(qn("w:shd"))
    if old is not None:
        properties.remove(old)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_paragraph_border(
    paragraph: Any,
    *,
    side: str = "bottom",
    color: str = RULE,
    size: int = 8,
    space: int = 4,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    existing = borders.find(qn(f"w:{side}"))
    if existing is not None:
        borders.remove(existing)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def _set_cell_margins(cell: Any, *, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)


def _prevent_row_split(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_table_borders(table: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), RULE)


def _set_table_layout_fixed(table: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _add_field(paragraph: Any, instruction: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    fallback_run = OxmlElement("w:r")
    fallback_text = OxmlElement("w:t")
    fallback_text.text = "1"
    fallback_run.append(fallback_text)
    field.append(fallback_run)
    paragraph._p.append(field)


def _add_hyperlink(paragraph: Any, label: str, target: str, *, size: float | None = None) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(target, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL_DARK)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    run_properties.extend((fonts, color, underline))
    if size is not None:
        half_points = OxmlElement("w:sz")
        half_points.set(qn("w:val"), str(int(round(size * 2))))
        run_properties.append(half_points)
        east_asian_size = OxmlElement("w:szCs")
        east_asian_size.set(qn("w:val"), str(int(round(size * 2))))
        run_properties.append(east_asian_size)
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _latex_to_readable(value: str) -> str:
    """Turn the small LaTeX subset used by the reports into readable linear math."""

    value = value.strip()
    value = re.sub(r"\\text\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\\operatorname\{([^{}]*)\}", r"\1", value)
    # Flatten braced indices before parsing fractions so v_{enc} is no longer
    # nested inside a denominator brace.
    value = re.sub(r"_\{([^{}]+)\}", r"_\1", value)
    value = re.sub(r"\^\{([^{}]+)\}", r"^\1", value)
    fraction = re.compile(r"\\frac\{([^{}]*)\}\s*\{([^{}]*)\}")
    previous = None
    while previous != value:
        previous = value
        value = fraction.sub(r"(\1)/(\2)", value)
    replacements = (
        (r"\operatorname", ""),
        (r"\qquad", "    "),
        (r"\propto", "∝"),
        (r"\lambda", "λ"),
        (r"\times", "×"),
        (r"\right", ""),
        (r"\left", ""),
        (r"\sqrt", "√"),
        (r"\quad", "  "),
        (r"\hat z", "ẑ"),
        (r"\log_2", " log₂"),
        (r"\rho", "ρ"),
        (r"\phi", "φ"),
        (r"\le", "≤"),
        (r"\ge", "≥"),
        (r"\in", "∈"),
        (r"\log", "log"),
        (r"\sum", "Σ"),
        (r"\,", " "),
    )
    for source, replacement in replacements:
        value = value.replace(source, replacement)
    value = value.replace("^2", "²").replace("^{-21}", "⁻²¹")
    value = value.replace("{", "").replace("}", "")
    unknown = sorted(set(re.findall(r"\\[A-Za-z]+", value)))
    if unknown:
        raise ValueError(f"지원하지 않는 LaTeX 명령: {', '.join(unknown)}")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _add_math_run(paragraph: Any, latex: str, *, size: float | None = None, color: str | None = None) -> None:
    run = paragraph.add_run(_latex_to_readable(latex))
    _set_font(
        run,
        size=size or 10.2,
        italic=True,
        color=color or NAVY,
        family="Cambria Math",
    )


def _add_inline_runs(paragraph: Any, text: str, *, size: float | None = None, color: str | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    text = _relativize_workspace_paths(text)
    token_re = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\\\(.+?\\\)|\[[^\]]+\]\([^)]+\))")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            _set_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_font(run, size=size, bold=True, color=NAVY if color is None else color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_font(run, size=(size or 10.2) - 0.5, color=TEAL_DARK, monospace=True)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:fill"), PALE_TEAL)
            run._element.get_or_add_rPr().append(shading)
        elif token.startswith(r"\("):
            _add_math_run(paragraph, token[2:-2], size=size, color=color)
        else:
            link_match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
            label, target = link_match.groups() if link_match else (token, "")
            _add_hyperlink(paragraph, label, target, size=size)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_font(run, size=size, color=color)


def _plain_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def _relativize_workspace_paths(text: str) -> str:
    """Keep machine-specific workspace prefixes out of reader-facing pages."""

    variants = {str(ROOT), str(ROOT).replace("\\", "/")}
    for prefix in sorted(variants, key=len, reverse=True):
        text = text.replace(prefix, ".")
    return text


def _report_code(markdown_path: Path) -> str:
    match = re.match(r"(0[1-4])", markdown_path.stem)
    return match.group(1) if match else "01"


def _extract_outline(
    lines: Sequence[str],
    is_table_separator: Callable[[str], bool],
) -> tuple[str, list[str], int, int]:
    title = "연구 보고서"
    headings: list[str] = []
    table_count = 0
    image_count = 0
    for index, line in enumerate(lines):
        stripped = line.strip()
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            value = _plain_markdown(heading.group(2))
            if level == 1 and title == "연구 보고서":
                title = value
            elif level == 2:
                headings.append(value)
        if re.fullmatch(r"!\[[^\]]*\]\([^)]+\)", stripped):
            image_count += 1
        if "|" in stripped and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            table_count += 1
    return title, headings, table_count, image_count


def _configure_document(document: Any, short_title: str) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = _rgb(INK)
    normal.paragraph_format.line_spacing = 1.38
    normal.paragraph_format.space_after = Pt(7)

    for name, size, color, before, after in (
        ("Heading 1", 22, NAVY, 20, 10),
        ("Heading 2", 16, NAVY, 17, 8),
        ("Heading 3", 12.5, TEAL_DARK, 12, 5),
        ("Heading 4", 10.8, NAVY_2, 9, 4),
    ):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    custom_styles = {
        "Report Body": (10.2, INK, 1.38, 7),
        "Report Callout": (9.7, INK, 1.3, 7),
        "Report Caption": (9.0, NAVY, 1.2, 3),
        "Report Source": (8.5, MUTED, 1.22, 9),
        "Report Code": (8.6, INK, 1.02, 8),
        "Report Formula": (11.0, NAVY, 1.1, 9),
        "Report Guide": (9.2, INK, 1.18, 4),
    }
    for name, (size, color, spacing, after) in custom_styles.items():
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = BODY_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = _rgb(color)
        style.paragraph_format.line_spacing = spacing
        style.paragraph_format.space_after = Pt(after)

    code_style = styles["Report Code"]
    code_style.font.name = MONO_FONT
    code_style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), MONO_FONT)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = BODY_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(10)
        style.font.color.rgb = _rgb(INK)
        style.paragraph_format.space_after = Pt(3.5)
        style.paragraph_format.line_spacing = 1.28

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.space_after = Pt(3)
    run = header_para.add_run(f"AIRTALKING REPLICATION  /  {short_title}")
    _set_font(run, size=7.5, bold=True, color=NAVY_2)
    _set_paragraph_border(header_para, color=TEAL, size=8, space=5)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.paragraph_format.space_before = Pt(2)
    run = footer_para.add_run("EVIDENCE-BACKED REPORT    ·    ")
    _set_font(run, size=7.2, color=MUTED)
    _add_field(footer_para, "PAGE")
    _set_paragraph_border(footer_para, side="top", color=RULE, size=4, space=4)

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _add_cover(
    document: Any,
    *,
    title: str,
    presentation: dict[str, str],
    headings: Sequence[str],
    table_count: int,
    image_count: int,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Inches, Pt

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(26)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(6)
    _set_shading(kicker._p, NAVY)
    run = kicker.add_run("  " + presentation["kicker"])
    _set_font(run, size=9.2, bold=True, color="8ED5D4")

    hero = document.add_paragraph(style="Heading 1")
    hero.paragraph_format.space_before = Pt(0)
    hero.paragraph_format.space_after = Pt(8)
    hero.paragraph_format.left_indent = Inches(0.12)
    hero.paragraph_format.right_indent = Inches(0.12)
    _set_shading(hero._p, NAVY)
    run = hero.add_run(title)
    _set_font(run, size=25, bold=True, color=WHITE)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(0)
    subtitle.paragraph_format.left_indent = Inches(0.12)
    subtitle.paragraph_format.right_indent = Inches(0.12)
    _set_shading(subtitle._p, NAVY)
    run = subtitle.add_run(presentation["subtitle"])
    _set_font(run, size=11.2, color="D8EAF0")

    chips = document.add_paragraph()
    chips.paragraph_format.space_before = Pt(9)
    chips.paragraph_format.space_after = Pt(16)
    chips.paragraph_format.left_indent = Inches(0.12)
    _set_shading(chips._p, NAVY)
    run = chips.add_run(presentation["chip"])
    _set_font(run, size=8.6, bold=True, color=WHITE)
    _set_paragraph_border(chips, color=TEAL, size=24, space=8)

    focus_label = document.add_paragraph()
    focus_label.paragraph_format.space_before = Pt(24)
    focus_label.paragraph_format.space_after = Pt(5)
    run = focus_label.add_run("READING FOCUS  /  빠르게 읽는 법")
    _set_font(run, size=9, bold=True, color=TEAL_DARK)

    focus = document.add_paragraph(style="Report Callout")
    focus.paragraph_format.left_indent = Inches(0.16)
    focus.paragraph_format.right_indent = Inches(0.08)
    focus.paragraph_format.space_after = Pt(14)
    _set_shading(focus._p, PALE_TEAL)
    _set_paragraph_border(focus, side="left", color=TEAL, size=18, space=8)
    _add_inline_runs(focus, presentation["focus"], size=10)

    stats = document.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.LEFT
    stats.paragraph_format.space_before = Pt(3)
    stats.paragraph_format.space_after = Pt(15)
    run = stats.add_run(
        f"{len(headings):02d}  SECTIONS     {table_count:02d}  TABLES     {image_count:02d}  FIGURES"
    )
    _set_font(run, size=9.2, bold=True, color=NAVY_2)

    source = document.add_paragraph(style="Report Source")
    source.paragraph_format.left_indent = Inches(0.16)
    source.paragraph_format.right_indent = Inches(0.08)
    _set_shading(source._p, PAPER)
    _set_paragraph_border(source, side="left", color=CORAL, size=14, space=7)
    run = source.add_run("그림 출처 원칙  ")
    _set_font(run, size=8, bold=True, color=CORAL)
    _add_inline_runs(
        source,
        "각 그림 아래에 원천 데이터·로컬 파일·생성 환경을 표기했습니다. "
        "Cityscapes 사진이 포함된 패널은 데이터셋 원출처와 본 연구의 모델 출력을 분리해 밝혔습니다.",
        size=7.8,
        color=MUTED,
    )

    break_run = document.add_paragraph().add_run()
    break_run.add_break(WD_BREAK.PAGE)


def _add_guide(
    document: Any,
    *,
    headings: Sequence[str],
    table_count: int,
    image_count: int,
) -> None:
    from docx.enum.text import WD_BREAK
    from docx.shared import Inches, Pt

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("문서 길잡이")
    _set_font(run, size=21, bold=True, color=NAVY)
    _set_paragraph_border(title, color=TEAL, size=18, space=8)

    intro = document.add_paragraph(style="Report Guide")
    intro.paragraph_format.space_after = Pt(10)
    _add_inline_runs(
        intro,
        f"본문은 **{len(headings)}개 큰 절**, 표 {table_count}개, 그림 {image_count}개로 구성됩니다. "
        "아래 순서를 훑은 뒤 필요한 절부터 읽어도 흐름이 끊기지 않도록 제목 체계를 정리했습니다.",
        size=9.5,
    )

    for index, heading in enumerate(headings, start=1):
        display_heading = re.sub(r"^\d+(?:\.\d+)*[.)]?\s+", "", heading)
        paragraph = document.add_paragraph(style="Report Guide")
        paragraph.paragraph_format.left_indent = Inches(0.10)
        paragraph.paragraph_format.first_line_indent = Inches(-0.02)
        paragraph.paragraph_format.space_before = Pt(0.5)
        paragraph.paragraph_format.space_after = Pt(2.2)
        paragraph.paragraph_format.keep_together = True
        _set_paragraph_border(paragraph, side="bottom", color="E6ECEF", size=3, space=2)
        run = paragraph.add_run(f"{index:02d}  ")
        _set_font(run, size=8.2, bold=True, color=TEAL)
        run = paragraph.add_run(display_heading)
        _set_font(run, size=9.1, color=INK)

    note = document.add_paragraph(style="Report Source")
    note.paragraph_format.space_before = Pt(9)
    _set_shading(note._p, PALE_BLUE)
    _set_paragraph_border(note, side="left", color=NAVY_2, size=12, space=6)
    _add_inline_runs(
        note,
        "표의 숫자는 짧게 비교하고, 해석 문단의 **가정·한계·증거 없음** 표기를 함께 확인하세요.",
        size=8.2,
    )

    break_run = document.add_paragraph().add_run()
    break_run.add_break(WD_BREAK.PAGE)


def _style_heading(paragraph: Any, level: int) -> None:
    from docx.shared import Inches

    if level == 2:
        paragraph.paragraph_format.left_indent = Inches(0.08)
        paragraph.paragraph_format.right_indent = Inches(0.02)
        _set_shading(paragraph._p, PALE_BLUE)
        _set_paragraph_border(paragraph, side="left", color=TEAL, size=24, space=8)
    elif level == 3:
        _set_paragraph_border(paragraph, color=RULE, size=4, space=4)


def _add_body_paragraph(document: Any, text: str) -> Any:
    paragraph = document.add_paragraph(style="Report Body")
    paragraph.paragraph_format.widow_control = True
    _add_inline_runs(paragraph, text)
    return paragraph


def _add_code_block(document: Any, code_lines: Sequence[str], language: str) -> None:
    from docx.shared import Inches, Pt

    if language:
        label = document.add_paragraph()
        label.paragraph_format.space_before = Pt(5)
        label.paragraph_format.space_after = Pt(2)
        label.paragraph_format.keep_with_next = True
        label.paragraph_format.keep_together = True
        run = label.add_run(language.upper())
        _set_font(run, size=8.0, bold=True, color=TEAL_DARK, monospace=True)

    paragraph = document.add_paragraph(style="Report Code")
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    _set_shading(paragraph._p, PAPER)
    _set_paragraph_border(paragraph, side="left", color=TEAL, size=16, space=7)
    run = paragraph.add_run(_relativize_workspace_paths("\n".join(code_lines)))
    _set_font(run, size=8.6, color=INK, monospace=True)


def _add_formula(document: Any, formula_lines: Sequence[str]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    paragraph = document.add_paragraph(style="Report Formula")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.right_indent = Inches(0.28)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.keep_together = True
    _set_shading(paragraph._p, PALE_BLUE)
    _set_paragraph_border(paragraph, side="left", color=NAVY_2, size=14, space=7)
    _add_math_run(paragraph, " ".join(line.strip() for line in formula_lines), size=11.0)


def _add_callout(document: Any, quote_lines: Sequence[str]) -> None:
    from docx.shared import Inches

    text = " ".join(line.strip().lstrip("> ") for line in quote_lines)
    paragraph = document.add_paragraph(style="Report Callout")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    _set_shading(paragraph._p, PALE_YELLOW)
    _set_paragraph_border(paragraph, side="left", color=CORAL, size=18, space=8)
    run = paragraph.add_run("NOTE  ")
    _set_font(run, size=8.2, bold=True, color=CORAL)
    _add_inline_runs(paragraph, text, size=9.7)


def _repository_display_path(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT).as_posix()
    except ValueError:
        return str(path.resolve())


def _figure_source_text(image_path: Path) -> str:
    display_path = _repository_display_path(image_path)
    kind = FIGURE_PROVENANCE.get(display_path)
    if kind == "cityscapes_panel":
        return (
            "출처: 원본 사진·정답 — Cityscapes 데이터셋(Cordts et al., CVPR 2016). "
            f"공식 인용: [Cityscapes Citation]({CITYSCAPES_CITATION_URL})\n"
            "모델 출력·패널 편집: 본 연구에서 직접 생성\n"
            "생성 환경: 연구자 본인 소유 PC(NVIDIA GeForce RTX 4060 Ti 탑재)\n"
            f"로컬 파일: `{display_path}`"
        )
    if kind == "airtalking_plot":
        return (
            "출처: 본 연구의 로컬 AirTalking 재현 시뮬레이션에서 직접 생성한 도표(외부 사진 아님)\n"
            "semantic codec 측정값: 연구자 본인 소유 PC의 NVIDIA GeForce RTX 4060 Ti 실행 산출물\n"
            "도표 렌더링: 저장된 로컬 시뮬레이션 결과\n"
            f"로컬 파일: `{display_path}`"
        )
    if kind == "adaptive_plot":
        return (
            "출처: 본 연구의 로컬 적응형 의미압축 시뮬레이션에서 직접 생성한 도표(외부 사진 아님)\n"
            "semantic codec profile: 연구자 본인 소유 PC의 NVIDIA GeForce RTX 4060 Ti 측정 산출물\n"
            "도표 렌더링: 저장된 로컬 시뮬레이션 결과\n"
            f"로컬 파일: `{display_path}`"
        )
    return (
        "출처 정보 미등록: 문서 입력으로 제공된 로컬 이미지입니다.\n"
        "제작자·원천 데이터·생성 장비를 확인한 뒤 배포해야 합니다.\n"
        f"로컬 파일: `{display_path}`"
    )


def _add_figure(
    document: Any,
    image_path: Path,
    alt: str,
    figure_number: int,
    warnings: list[str],
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(9)
    picture_paragraph.paragraph_format.space_after = Pt(4)
    picture_paragraph.paragraph_format.keep_with_next = True
    run = picture_paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.55))
    max_height = Inches(6.65)
    if shape.height > max_height:
        ratio = max_height / shape.height
        shape.height = max_height
        shape.width = int(shape.width * ratio)
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", f"그림 {figure_number}. {alt}")

    caption = document.add_paragraph(style="Report Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run(f"그림 {figure_number:02d}  |  ")
    _set_font(run, size=8.8, bold=True, color=TEAL_DARK)
    run = caption.add_run(alt)
    _set_font(run, size=9.0, bold=True, color=NAVY)

    source = document.add_paragraph(style="Report Source")
    source.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source.paragraph_format.left_indent = Inches(0.18)
    source.paragraph_format.right_indent = Inches(0.18)
    source.paragraph_format.keep_together = True
    _set_shading(source._p, PAPER)
    _set_paragraph_border(source, side="left", color=TEAL, size=10, space=6)
    if _repository_display_path(image_path) not in FIGURE_PROVENANCE:
        warnings.append(f"DOCX 그림 출처 정보 미등록: {image_path}")
    _add_inline_runs(source, _figure_source_text(image_path), size=8.4, color=MUTED)


def _add_end_mark(document: Any, presentation: dict[str, str]) -> None:
    from docx.shared import Pt

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(24)
    closing.paragraph_format.space_after = Pt(4)
    closing.paragraph_format.keep_together = True
    _set_shading(closing._p, NAVY)
    _set_paragraph_border(closing, side="top", color=TEAL, size=18, space=7)
    run = closing.add_run(f"  END OF REPORT   /   {presentation['short']}")
    _set_font(run, size=9.0, bold=True, color=WHITE)

    note = document.add_paragraph(style="Report Source")
    note.paragraph_format.space_after = Pt(0)
    _set_shading(note._p, PAPER)
    _add_inline_runs(
        note,
        "문서 끝 · 수치와 해석의 근거는 최종 Markdown과 manifest에 보존되어 있습니다.",
        size=8.4,
        color=MUTED,
    )


def _rewrite_extended_properties(docx_path: Path) -> None:
    """Remove stale template statistics that python-docx otherwise preserves."""

    app_part = "docProps/app.xml"
    namespace = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    removable = {
        "Pages",
        "Words",
        "Characters",
        "CharactersWithSpaces",
        "Lines",
        "Paragraphs",
        "TotalTime",
    }
    temporary = docx_path.with_suffix(docx_path.suffix + ".props.tmp")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED
    ) as destination:
        for info in source.infolist():
            payload = source.read(info.filename)
            if info.filename == app_part:
                root = ElementTree.fromstring(payload)
                for child in list(root):
                    if child.tag.rsplit("}", 1)[-1] in removable:
                        root.remove(child)
                application = root.find(f"{{{namespace}}}Application")
                if application is None:
                    application = ElementTree.SubElement(root, f"{{{namespace}}}Application")
                application.text = "AirTalking DOCX Renderer"
                app_version = root.find(f"{{{namespace}}}AppVersion")
                if app_version is not None:
                    app_version.text = "1.0"
                payload = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
            destination.writestr(info, payload)
    temporary.replace(docx_path)


def _table_font_size(column_count: int) -> float:
    if column_count <= 2:
        return 8.8
    if column_count <= 4:
        return 8.3
    if column_count <= 6:
        return 8.0
    if column_count <= 8:
        return 7.7
    return 7.5


def _column_weights(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> list[float]:
    weights: list[float] = []
    for index, header in enumerate(headers):
        lengths = [len(_plain_markdown(header))]
        lengths.extend(len(_plain_markdown(row[index])) for row in rows)
        longest = max(lengths, default=1)
        # Keep path/provenance columns useful without allowing one cell to consume the page.
        weights.append(float(max(5, min(longest, 26))))
    total = sum(weights) or 1.0
    minimum = 0.075 if len(headers) >= 8 else 0.12 if len(headers) >= 5 else 0.18
    normalized = [max(minimum, value / total) for value in weights]
    norm_total = sum(normalized)
    return [value / norm_total for value in normalized]


def _add_table(document: Any, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_layout_fixed(table)
    _set_table_borders(table)
    font_size = _table_font_size(len(headers))
    widths = _column_weights(headers, rows)
    # A4 width minus the two 0.72-inch margins is about 6.83 inches.
    available_width = 6.72
    for column, weight in enumerate(widths):
        table.columns[column].width = Inches(available_width * weight)

    header_row = table.rows[0]
    _set_repeat_table_header(header_row)
    _prevent_row_split(header_row)
    for column, value in enumerate(headers):
        cell = header_row.cells[column]
        cell.width = Inches(available_width * widths[column])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_shading(cell._tc, NAVY)
        _set_cell_margins(cell, top=100, start=90, bottom=100, end=90)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        _add_inline_runs(paragraph, value, size=font_size, color=WHITE)
        for run in paragraph.runs:
            run.bold = True

    for row_index, values in enumerate(rows, start=1):
        row = table.add_row()
        _prevent_row_split(row)
        fill = WHITE if row_index % 2 else PAPER
        for column, value in enumerate(values):
            cell = row.cells[column]
            cell.width = Inches(available_width * widths[column])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_shading(cell._tc, fill)
            _set_cell_margins(cell, top=80, start=85, bottom=80, end=85)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.06
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if re.fullmatch(r"[-+−]?\d[\d.,%+−\-–— /:]*", _plain_markdown(value))
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            _add_inline_runs(paragraph, value.replace("<br>", "\n"), size=font_size)

def build_docx(
    markdown_path: Path,
    docx_path: Path,
    warnings: list[str],
    *,
    split_md_row: Callable[[str], list[str]],
    is_table_separator: Callable[[str], bool],
    generated_at: datetime | None = None,
) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Inches, Pt

    lines = markdown_path.read_text(encoding="utf-8-sig").splitlines()
    title, headings, table_count, image_count = _extract_outline(lines, is_table_separator)
    code = _report_code(markdown_path)
    presentation = REPORT_PRESENTATION.get(code, REPORT_PRESENTATION["01"])

    document = Document()
    _configure_document(document, presentation["short"])
    _add_cover(
        document,
        title=title,
        presentation=presentation,
        headings=headings,
        table_count=table_count,
        image_count=image_count,
    )
    _add_guide(
        document,
        headings=headings,
        table_count=table_count,
        image_count=image_count,
    )

    index = 0
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    paragraph_lines: list[str] = []
    figure_number = 0
    skipped_title = False

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        _add_body_paragraph(document, " ".join(line.strip() for line in paragraph_lines))
        paragraph_lines.clear()

    def flush_code() -> None:
        if not code_lines:
            return
        _add_code_block(document, code_lines, code_language)
        code_lines.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                flush_code()
                in_code = False
                code_language = ""
            else:
                in_code = True
                code_language = stripped[3:].strip()
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped == r"\[":
            flush_paragraph()
            formula_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != r"\]":
                formula_lines.append(lines[index])
                index += 1
            if index < len(lines) and lines[index].strip() == r"\]":
                index += 1
            _add_formula(document, formula_lines)
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            flush_paragraph()
            alt, raw_path = image_match.groups()
            raw_path = raw_path.strip().strip("<>")
            image_path = Path(raw_path)
            if not image_path.is_absolute():
                image_path = (markdown_path.parent / image_path).resolve()
            if image_path.is_file():
                try:
                    figure_number += 1
                    _add_figure(document, image_path, alt, figure_number, warnings)
                except Exception as exc:  # python-docx/Pillow raise several concrete types
                    warnings.append(f"DOCX 그림 삽입 실패: {image_path}: {exc}")
                    paragraph = _add_body_paragraph(document, f"[그림 삽입 실패: {alt} — {image_path}]")
                    _set_shading(paragraph._p, PALE_CORAL)
            else:
                warnings.append(f"DOCX 로컬 그림 경로가 오래되었거나 없습니다: {image_path}")
                paragraph = _add_body_paragraph(document, f"[그림 증거 없음: {alt} — {image_path}]")
                _set_shading(paragraph._p, PALE_CORAL)
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            raw_level = len(heading.group(1))
            if raw_level == 1 and not skipped_title:
                skipped_title = True
                index += 1
                continue
            level = min(raw_level, 4)
            paragraph = document.add_heading(level=level)
            _add_inline_runs(paragraph, heading.group(2))
            _style_heading(paragraph, level)
            index += 1
            continue

        if "|" in stripped and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            flush_paragraph()
            headers = split_md_row(line)
            table_rows: list[list[str]] = []
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                row = split_md_row(lines[index])
                if len(row) != len(headers):
                    break
                table_rows.append(row)
                index += 1
            _add_table(document, headers, table_rows)
            continue

        bullet = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        numbered = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if bullet or numbered:
            flush_paragraph()
            match = bullet or numbered
            style = "List Bullet" if bullet else "List Number"
            paragraph = document.add_paragraph(style=style)
            indent_level = min(len(match.group(1).expandtabs(4)) // 2, 4)
            paragraph.paragraph_format.left_indent = Inches(0.22 + 0.22 * indent_level)
            paragraph.paragraph_format.first_line_indent = Inches(-0.14)
            _add_inline_runs(paragraph, match.group(2), size=10)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index])
                index += 1
            _add_callout(document, quote_lines)
            continue

        if re.fullmatch(r"[-*_]{3,}", stripped):
            flush_paragraph()
            rule = document.add_paragraph()
            rule.paragraph_format.space_before = Pt(5)
            rule.paragraph_format.space_after = Pt(7)
            _set_paragraph_border(rule, color=TEAL, size=9, space=2)
            index += 1
            continue

        paragraph_lines.append(line)
        index += 1

    flush_paragraph()
    if in_code:
        flush_code()

    _add_end_mark(document, presentation)

    document.core_properties.title = title
    document.core_properties.subject = presentation["subtitle"]
    document.core_properties.author = "AirTalking Replication Experiment"
    document.core_properties.last_modified_by = "AirTalking DOCX report renderer"
    build_time = generated_at or datetime.now(timezone.utc)
    document.core_properties.created = build_time
    document.core_properties.modified = build_time
    document.core_properties.revision = 1
    document.core_properties.keywords = "AirTalking, semantic communication, RTX 4060 Ti, reproducible research"
    document.core_properties.comments = (
        "Figures carry explicit source notes. Cityscapes-origin imagery is cited separately from locally generated model output."
    )
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)
    _rewrite_extended_properties(docx_path)
