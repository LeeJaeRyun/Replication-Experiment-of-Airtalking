from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]

REPRO = ROOT / "studies" / "airtalking_reproduction"
ADAPT = ROOT / "studies" / "adaptive_semantic_compression"
NEURAL = ROOT / "studies" / "neural_encoder_decoder"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(96, 106, 116)
INK = RGBColor(20, 20, 20)
FILL = "F2F4F7"
CALLOUT = "F4F6F9"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def f(value: str | float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def set_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top: int = 90, start: int = 130, bottom: int = 90, end: int = 130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    dxa_widths = [int(round(width * 1440)) for width in widths]
    for dxa in dxa_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(dxa))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa_widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)


def new_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = title
    set_font(header.runs[0], 9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "AirTalking 연구 정리"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.runs[0], 9, color=MUTED)

    p = doc.add_paragraph()
    spacing(p, after=4)
    r = p.add_run(title)
    set_font(r, 22, True, INK)
    p = doc.add_paragraph()
    spacing(p, after=14)
    r = p.add_run(subtitle)
    set_font(r, 12, color=MUTED)
    return doc


def para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    spacing(p)
    r = p.add_run(text)
    set_font(r, 11, bold=bold)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    spacing(p, after=4, line=1.167)
    r = p.add_run(text)
    set_font(r, 11)


def callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, CALLOUT)
    p = cell.paragraphs[0]
    spacing(p, after=0)
    r = p.add_run(text)
    set_font(r, 11, True, DARK)
    table_geometry(table, [6.5])
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, after=0)
        r = p.add_run(header)
        set_font(r, 10, True)
    for values in rows:
        row = table.add_row()
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            spacing(p, after=0)
            r = p.add_run(text)
            set_font(r, 10)
    table_geometry(table, widths)
    doc.add_paragraph()


def lookup(rows: list[dict[str, str]]) -> dict[tuple[str, int, str], dict[str, str]]:
    return {(row["mode"], int(float(row["area"])), row["policy"]): row for row in rows}


def build_reproduction_doc() -> Path:
    repro_summary = lookup(read_csv(REPRO / "results" / "airtalking_cityscapes_calibrated_final_p012" / "summary_metrics.csv"))
    neural_summary = lookup(read_csv(NEURAL / "results" / "airtalking_neural_encoder_decoder_timed" / "summary_metrics.csv"))
    neural_profile = json.loads((NEURAL / "results" / "paperlike_timed_latent20" / "airtalking_semantic_summary.json").read_text(encoding="utf-8"))

    doc = new_doc("AirTalking 재현 실험 쉬운 보고서", "논문 재현 + neural encoder/decoder 반영 위치")
    callout(doc, "결론: 이 폴더는 원 논문을 따라 만든 재현 실험이다. 이제 neural encoder/decoder 결과도 같이 참조하도록 정리했다.")
    doc.add_heading("1. 이 폴더는 무엇인가", level=1)
    para(doc, "AirTalking 논문에서 나온 UAV D2D semantic communication 실험을 공개 데이터로 최대한 따라 한 폴더다.")
    bullet(doc, "Cityscapes train/val 데이터로 semantic payload 크기를 측정했다.")
    bullet(doc, "논문에 공개된 파라미터와 명시되지 않은 가정값을 분리해서 시뮬레이터에 넣었다.")
    bullet(doc, "논문 그래프와 같은 방향인지 확인하기 위해 결과를 CSV와 PNG로 저장했다.")

    doc.add_heading("2. encoder/decoder는 어디에 반영됐나", level=1)
    para(doc, "기존 재현 실험은 논문 Table III의 encoding/decoding 속도와 Cityscapes 기반 payload ratio를 썼다. 여기에 trained neural encoder/decoder 결과를 넣을 수 있도록 시뮬레이터 코드를 수정했다.")
    add_table(
        doc,
        ["항목", "값"],
        [
            ["neural rho_c", f(neural_profile["rho_c_feature_uncompressed_mean"], 5)],
            ["pixel accuracy", f(neural_profile["pixel_accuracy_best"], 4)],
            ["mIoU", f(neural_profile["semantic_quality_miou_best"], 4)],
            ["encode/decode", f"{f(neural_profile['timing']['encode_ms_median'], 2)} ms / {f(neural_profile['timing']['decode_ms_median'], 2)} ms"],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("3. 300m 결과 예시", level=1)
    rows = []
    for policy in ["LinUCB", "SA", "Greedy", "MCTS"]:
        base = repro_summary[("semantic", 300, policy)]
        neural = neural_summary[("semantic", 300, policy)]
        rows.append([policy, f(base["finished"], 1), f(neural["finished"], 1), f(base["avg_time"], 2), f(neural["avg_time"], 2)])
    add_table(doc, ["정책", "기존 semantic 완료", "neural 반영 완료", "기존 평균시간", "neural 평균시간"], rows, [1.05, 1.35, 1.35, 1.35, 1.40])

    doc.add_heading("4. 읽을 때 주의할 점", level=1)
    bullet(doc, "완전한 원 논문 복제는 아니다. 원 논문 코드와 원본 raw plotting data가 공개되지 않았기 때문이다.")
    bullet(doc, "neural encoder/decoder는 실제 학습했지만, 작은 CPU 실험용 모델이다.")
    bullet(doc, "따라서 재현 실험은 논문값 근사와 후속 연구 검증용 기반으로 보는 것이 맞다.")

    path = REPRO / "reports" / "AirTalking_Reproduction_Easy_Final_KR.docx"
    doc.save(path)
    return path


def build_adaptive_doc() -> Path:
    adaptive = lookup(read_csv(ADAPT / "results" / "full_adaptive_results" / "summary_metrics.csv"))
    neural_profile = json.loads((NEURAL / "results" / "paperlike_timed_latent20" / "airtalking_semantic_summary.json").read_text(encoding="utf-8"))

    doc = new_doc("Adaptive Semantic Compression 쉬운 보고서", "상황별 압축률 변경 실험 + neural encoder/decoder 반영")
    callout(doc, "결론: 이 폴더는 채널 상태가 나쁠 때는 더 세게 압축하고, 좋을 때는 품질을 더 살리는 방식이 효과가 있는지 본 실험이다.")
    doc.add_heading("1. 이 폴더는 무엇인가", level=1)
    para(doc, "원래 재현 실험은 semantic compression을 거의 고정값처럼 썼다. 이 폴더는 링크 상태에 따라 압축률을 바꾸는 실험이다.")
    bullet(doc, "링크가 나쁘면 payload를 줄이기 위해 더 작은 semantic 표현을 선택한다.")
    bullet(doc, "링크가 좋으면 품질을 위해 더 큰 semantic 표현을 선택할 수 있다.")
    bullet(doc, "결과는 fixed paper-like 방식과 adaptive 방식을 비교한다.")

    doc.add_heading("2. encoder/decoder는 어떻게 반영됐나", level=1)
    para(doc, "현재 학습된 neural encoder/decoder는 paper-like 압축률 하나만 있다. 그래서 adaptive 전체를 모두 neural network로 바꾼 것은 아니고, paper_like 단계의 기준값으로 neural 결과를 연결했다.")
    add_table(
        doc,
        ["항목", "반영 방식"],
        [
            ["rho_c", f"trained encoder/decoder 값 {f(neural_profile['rho_c_feature_uncompressed_mean'], 5)}를 paper_like 기준값으로 사용"],
            ["encode/decode 시간", f"{f(neural_profile['timing']['encode_ms_median'], 2)} ms / {f(neural_profile['timing']['decode_ms_median'], 2)} ms를 기록"],
            ["품질", f"기본 모드는 record_only. 현재 mIoU {f(neural_profile['semantic_quality_miou_best'], 4)}는 별도로 기록하고, 다른 압축 단계는 아직 proxy quality 사용"],
        ],
        [1.4, 5.1],
    )

    doc.add_heading("3. 300m 결과 예시", level=1)
    rows = []
    for policy in ["LinUCB", "SA", "Greedy", "MCTS"]:
        fixed = adaptive[("fixed_paper_like", 300, policy)]
        ada = adaptive[("adaptive_semantic", 300, policy)]
        rows.append([policy, f(fixed["finished"], 1), f(ada["finished"], 1), f(fixed["avg_time"], 2), f(ada["avg_time"], 2)])
    add_table(doc, ["정책", "fixed 완료", "adaptive 완료", "fixed 평균시간", "adaptive 평균시간"], rows, [1.1, 1.35, 1.35, 1.35, 1.35])

    doc.add_heading("4. 한계", level=1)
    bullet(doc, "진짜 완성형 adaptive neural codec이 되려면 여러 압축률별 encoder/decoder를 각각 학습해야 한다.")
    bullet(doc, "현재는 paper_like 압축 단계만 실제 neural encoder/decoder 결과로 고정했고, 나머지 단계는 Cityscapes label proxy다.")
    bullet(doc, "따라서 이 폴더는 '상황별 압축률 정책이 의미 있는가'를 확인하는 후속 실험이다.")

    path = ADAPT / "reports" / "Adaptive_Semantic_Compression_Easy_Final_KR.docx"
    doc.save(path)
    return path


def build_neural_doc() -> Path:
    summary = json.loads((NEURAL / "results" / "paperlike_timed_latent20" / "result_summary.json").read_text(encoding="utf-8"))
    sim = lookup(read_csv(NEURAL / "results" / "airtalking_neural_encoder_decoder_timed" / "summary_metrics.csv"))

    doc = new_doc("Neural Encoder/Decoder 쉬운 보고서", "Cityscapes로 학습한 실제 semantic encoder/decoder")
    callout(doc, "결론: 이 폴더는 논문에 나온 semantic encoder/decoder를 작은 딥러닝 모델로 직접 구현한 실험이다.")
    doc.add_heading("1. 이 폴더는 무엇인가", level=1)
    para(doc, "RGB 이미지를 encoder에 넣으면 작은 latent feature가 나오고, decoder가 이를 semantic segmentation map으로 복원한다.")
    bullet(doc, "입력: Cityscapes RGB image")
    bullet(doc, "출력: semantic segmentation map")
    bullet(doc, "목표: 논문 Table III의 rho_c=0.104와 비슷한 압축률 만들기")

    doc.add_heading("2. 모델 결과", level=1)
    model = summary["model"]
    metrics = summary["best_metrics"]
    timing = summary["timing"]
    add_table(
        doc,
        ["항목", "값"],
        [
            ["모델", f"{model['name']}, latent {model['latent_channels']} channels, 1/{model['downsample_factor']} 해상도"],
            ["rho_c", f(model["payload_ratio"], 5)],
            ["pixel accuracy", f(metrics["val_pixel_accuracy"], 4)],
            ["mIoU", f(metrics["val_mean_iou"], 4)],
            ["encode/decode/full", f"{f(timing['encode_ms_median'], 2)} / {f(timing['decode_ms_median'], 2)} / {f(timing['full_ms_median'], 2)} ms"],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("3. AirTalking에 넣었을 때", level=1)
    rows = []
    for policy in ["LinUCB", "SA", "Greedy", "MCTS"]:
        sem = sim[("semantic", 300, policy)]
        ns = sim[("nonsemantic", 300, policy)]
        rows.append([policy, f(sem["finished"], 1), f(ns["finished"], 1), f(sem["avg_time"], 2), f(ns["avg_time"], 2)])
    add_table(doc, ["정책", "semantic 완료", "nonsemantic 완료", "semantic 평균시간", "nonsemantic 평균시간"], rows, [1.05, 1.35, 1.35, 1.35, 1.40])

    doc.add_heading("4. 한계", level=1)
    bullet(doc, "논문 원본 encoder/decoder 코드와 weight는 공개되지 않았다.")
    bullet(doc, "현재 모델은 작은 CPU 학습 모델이라 mIoU가 높지는 않다.")
    bullet(doc, "그래도 실제 neural encoder/decoder를 만들고 rho_c를 논문값에 가깝게 맞춘 점이 핵심이다.")

    path = NEURAL / "reports" / "Neural_Encoder_Decoder_Easy_Final_KR.docx"
    doc.save(path)
    return path


def main() -> None:
    for reports in [REPRO / "reports", ADAPT / "reports", NEURAL / "reports"]:
        reports.mkdir(parents=True, exist_ok=True)
    paths = [build_reproduction_doc(), build_adaptive_doc(), build_neural_doc()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
